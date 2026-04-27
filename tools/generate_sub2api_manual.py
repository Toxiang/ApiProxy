from __future__ import annotations

import os
import re
import textwrap
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
IMG_DIR = OUT_DIR / "images"
DOCX_PATH = OUT_DIR / "Sub2API_前后端功能与代码说明书.docx"


def read_text(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def find_font() -> str | None:
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return item
    return None


FONT_PATH = find_font()


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH:
        try:
            return ImageFont.truetype(FONT_PATH, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def wrap_draw_text(draw: ImageDraw.ImageDraw, xy, text, width, fill, font_obj, line_gap=8):
    x, y = xy
    line = ""
    lines: list[str] = []
    for ch in text:
        test = line + ch
        if draw.textbbox((0, 0), test, font=font_obj)[2] <= width:
            line = test
        else:
            if line:
                lines.append(line)
            line = ch
    if line:
        lines.append(line)
    for ln in lines:
        draw.text((x, y), ln, fill=fill, font=font_obj)
        y += font_obj.size + line_gap
    return y


def rounded_box(draw, xy, fill, outline, radius=24, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill=(80, 90, 110), width=4):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    dx, dy = ex - sx, ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx > 0 else -1
        pts = [(ex, ey), (ex - 14 * sign, ey - 8), (ex - 14 * sign, ey + 8)]
    else:
        sign = 1 if dy > 0 else -1
        pts = [(ex, ey), (ex - 8, ey - 14 * sign), (ex + 8, ey - 14 * sign)]
    draw.polygon(pts, fill=fill)


def make_architecture_image(path: Path):
    img = Image.new("RGB", (1600, 980), "#f6f8fb")
    d = ImageDraw.Draw(img)
    title_font = font(42, True)
    h_font = font(28, True)
    body_font = font(23)
    small_font = font(20)

    d.text((70, 45), "Sub2API 总体架构图", fill="#172033", font=title_font)
    d.text((72, 103), "浏览器后台、API 客户端、Go 网关、数据库缓存、上游 AI 服务之间的关系", fill="#637083", font=small_font)

    boxes = {
        "client": (70, 210, 330, 405, "用户/客户端", "浏览器后台\nClaude Code / Codex\nOpenAI/Gemini SDK"),
        "frontend": (460, 150, 760, 335, "Vue 前端", "页面路由\nPinia 状态\nAxios API 调用"),
        "gateway": (460, 470, 760, 705, "Go 后端网关", "Gin 路由\n鉴权中间件\n业务 Service\n请求转发与计费"),
        "db": (915, 175, 1205, 355, "PostgreSQL", "用户、API Key\n账号、分组\n订单、用量日志"),
        "redis": (915, 450, 1205, 620, "Redis", "Token/鉴权缓存\n限流计数\n并发与调度缓存"),
        "upstream": (1270, 260, 1530, 525, "上游 AI", "Anthropic\nOpenAI/Codex\nGemini\nAntigravity"),
        "payment": (1270, 650, 1530, 820, "支付/外部服务", "EasyPay\n支付宝/微信\nStripe\nS3/备份"),
    }
    palette = {
        "client": ("#ffffff", "#90a4c8"),
        "frontend": ("#e8f8ef", "#39a76d"),
        "gateway": ("#eef4ff", "#4267d5"),
        "db": ("#fff5e6", "#e49b32"),
        "redis": ("#ffecec", "#d44d4d"),
        "upstream": ("#f1edff", "#7b61d1"),
        "payment": ("#eaf9fb", "#25a2b8"),
    }
    for key, (x1, y1, x2, y2, title, body) in boxes.items():
        fill, outline = palette[key]
        rounded_box(d, (x1, y1, x2, y2), fill, outline, radius=26, width=3)
        d.text((x1 + 28, y1 + 24), title, fill="#172033", font=h_font)
        wrap_draw_text(d, (x1 + 28, y1 + 72), body, x2 - x1 - 56, "#39465a", body_font)

    arrow(d, (330, 295), (460, 245))
    arrow(d, (610, 335), (610, 470))
    arrow(d, (760, 565), (915, 265))
    arrow(d, (760, 590), (915, 535))
    arrow(d, (760, 560), (1270, 390))
    arrow(d, (760, 635), (1270, 735))
    arrow(d, (1270, 420), (760, 580), fill=(120, 95, 180))

    d.text((88, 850), "核心理解：前端负责“看得见的管理界面”，后端负责“鉴权、调度、转发、计费、记录”，数据库和 Redis 保存状态，上游 AI 账号提供真正模型能力。", fill="#26364d", font=body_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_gateway_flow_image(path: Path):
    img = Image.new("RGB", (1600, 900), "#fbfcfe")
    d = ImageDraw.Draw(img)
    title_font = font(40, True)
    h_font = font(25, True)
    body_font = font(21)
    d.text((70, 45), "一次模型请求在后端里的流转", fill="#172033", font=title_font)
    steps = [
        ("1. 客户端请求", "携带 sk- API Key\n调用 /v1/messages、/responses 等"),
        ("2. 中间件检查", "请求大小、Request ID\nAPI Key 鉴权、分组检查"),
        ("3. 调度账号", "根据分组、平台、模型\n筛选可用上游账号"),
        ("4. 转换协议", "Claude/OpenAI/Gemini\n之间做兼容转换"),
        ("5. 转发上游", "请求 Anthropic/OpenAI\nGemini/Antigravity"),
        ("6. 计费记录", "统计 Token/图片/时长\n写入 UsageLog 与缓存"),
    ]
    x = 65
    y = 220
    w = 225
    h = 220
    for idx, (title, body) in enumerate(steps):
        x1 = x + idx * 250
        rounded_box(d, (x1, y, x1 + w, y + h), "#eef4ff" if idx % 2 == 0 else "#eaf9f2", "#5672c8" if idx % 2 == 0 else "#40a36d", 24, 3)
        d.text((x1 + 20, y + 24), title, fill="#172033", font=h_font)
        wrap_draw_text(d, (x1 + 20, y + 75), body, w - 40, "#3e4a5f", body_font)
        if idx < len(steps) - 1:
            arrow(d, (x1 + w + 8, y + h // 2), (x1 + 250 - 12, y + h // 2))
    rounded_box(d, (130, 560, 1470, 760), "#fff8e8", "#e3a943", 28, 3)
    d.text((170, 595), "失败处理与观测", fill="#172033", font=h_font)
    text = "如果上游账号限流、过载、超时或返回错误，后端会记录运维错误、尝试 failover、更新账号临时不可调度状态，并在 Ops 看板里展示 QPS、延迟、错误分布和系统日志。"
    wrap_draw_text(d, (170, 645), text, 1260, "#3e4a5f", body_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_frontend_map_image(path: Path):
    img = Image.new("RGB", (1600, 950), "#f7f9fc")
    d = ImageDraw.Draw(img)
    title_font = font(40, True)
    h_font = font(26, True)
    body_font = font(21)
    d.text((70, 45), "前端功能地图", fill="#172033", font=title_font)
    groups = [
        ("初始化/公开", "/setup\n/home\n/login /register\n回调与找回密码", "#fff5e6", "#dc9235"),
        ("用户中心", "/dashboard\n/keys\n/usage\n/redeem\n/profile\n/subscriptions", "#e8f8ef", "#39a76d"),
        ("管理员后台", "/admin/dashboard\n/admin/accounts\n/admin/users\n/admin/groups\n/admin/settings\n/admin/ops", "#eef4ff", "#4267d5"),
        ("支付订单", "/purchase\n/orders\n/payment/*\n/admin/payment/*", "#f1edff", "#7b61d1"),
        ("监控与扩展", "/monitor\n/admin/channels/*\n/custom/:id\n公告/自定义菜单", "#eaf9fb", "#25a2b8"),
    ]
    for i, (title, body, fill, outline) in enumerate(groups):
        col = i % 3
        row = i // 3
        x1 = 85 + col * 500
        y1 = 180 + row * 330
        rounded_box(d, (x1, y1, x1 + 430, y1 + 250), fill, outline, 26, 3)
        d.text((x1 + 25, y1 + 24), title, fill="#172033", font=h_font)
        wrap_draw_text(d, (x1 + 25, y1 + 74), body, 370, "#3d4a5f", body_font)
    d.text((90, 820), "代码关系：views 是页面，components 是可复用积木，stores 保存登录/设置/支付等全局状态，api 封装和后端交互的 Axios 请求。", fill="#26364d", font=body_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_domain_model_image(path: Path):
    img = Image.new("RGB", (1600, 980), "#fbfcff")
    d = ImageDraw.Draw(img)
    title_font = font(40, True)
    h_font = font(25, True)
    body_font = font(20)
    d.text((70, 45), "核心数据关系图", fill="#172033", font=title_font)
    d.text((72, 96), "理解 User、APIKey、Group、Account、UsageLog、PaymentOrder 的关系，是排查权限、调度和计费问题的关键。", fill="#637083", font=body_font)

    boxes = [
        ("User\n用户", 80, 250, "#eef4ff", "#4267d5", "登录、余额、角色\n并发、2FA"),
        ("APIKey\n访问密钥", 390, 250, "#e8f8ef", "#39a76d", "sk-xxx\n额度/IP/状态"),
        ("Group\n分组/套餐池", 700, 250, "#fff5e6", "#e49b32", "平台、倍率\n模型和限制"),
        ("Account\n上游账号", 1010, 250, "#f1edff", "#7b61d1", "OAuth/API Key\n状态、代理、并发"),
        ("UsageLog\n用量日志", 700, 610, "#eaf9fb", "#25a2b8", "Token、成本\n耗时、模型"),
        ("PaymentOrder\n支付订单", 80, 610, "#ffecec", "#d44d4d", "金额、通道\n履约、退款"),
        ("Subscription\n用户订阅", 390, 610, "#f7f1e8", "#b98a4a", "有效期、额度\n每日/周/月窗口"),
    ]
    rects = {}
    for title, x, y, fill, outline, body in boxes:
        rect = (x, y, x + 240, y + 180)
        rects[title.split("\n")[0]] = rect
        rounded_box(d, rect, fill, outline, 24, 3)
        d.text((x + 24, y + 20), title, fill="#172033", font=h_font)
        wrap_draw_text(d, (x + 24, y + 88), body, 190, "#3d4a5f", body_font)

    arrow(d, (320, 340), (390, 340))
    d.text((322, 300), "创建/持有", fill="#536071", font=body_font)
    arrow(d, (630, 340), (700, 340))
    d.text((626, 300), "绑定", fill="#536071", font=body_font)
    arrow(d, (940, 340), (1010, 340))
    d.text((935, 300), "调度账号池", fill="#536071", font=body_font)
    arrow(d, (820, 430), (820, 610))
    d.text((850, 505), "请求后写日志", fill="#536071", font=body_font)
    arrow(d, (320, 700), (390, 700))
    d.text((312, 660), "购买后履约", fill="#536071", font=body_font)
    arrow(d, (510, 610), (730, 430))
    d.text((545, 525), "赋予分组权益", fill="#536071", font=body_font)

    rounded_box(d, (80, 835, 1480, 910), "#f6f8fb", "#c8d2e3", 18, 2)
    d.text((110, 855), "排障顺序：先查 APIKey 是否有效，再查 Group 是否可用，再查 Account 是否可调度，最后看 UsageLog / Ops 错误确认上游结果。", fill="#26364d", font=body_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_deployment_image(path: Path):
    img = Image.new("RGB", (1600, 900), "#f8fafc")
    d = ImageDraw.Draw(img)
    title_font = font(40, True)
    h_font = font(25, True)
    body_font = font(20)
    d.text((70, 45), "部署形态对比图", fill="#172033", font=title_font)
    lanes = [
        ("Docker Compose 推荐路径", 90, "#eef4ff", "#4267d5", [
            ("apiproxy 容器", "Go 后端 + 内嵌前端"),
            ("postgres 容器", "业务数据库"),
            ("redis 容器", "缓存/限流/调度"),
            ("volume/data", "config、DB、Redis 数据"),
        ]),
        ("裸机 systemd 路径", 470, "#e8f8ef", "#39a76d", [
            ("/opt/sub2api/sub2api", "主程序"),
            ("sub2api.service", "systemd 常驻"),
            ("外部 PostgreSQL/Redis", "自行安装维护"),
            ("/etc/sub2api", "配置目录"),
        ]),
        ("数据管理 Agent", 850, "#fff5e6", "#e49b32", [
            ("datamanagementd", "宿主机守护进程"),
            ("/tmp/*.sock", "容器/主程序通信"),
            ("S3/本地备份", "备份恢复任务"),
            ("/var/lib/sub2api", "agent 数据库"),
        ]),
    ]
    for title, x, fill, outline, items in lanes:
        rounded_box(d, (x, 150, x + 420, 760), fill, outline, 28, 3)
        d.text((x + 28, 185), title, fill="#172033", font=h_font)
        y = 260
        for name, desc in items:
            rounded_box(d, (x + 35, y, x + 385, y + 92), "#ffffff", outline, 18, 2)
            d.text((x + 58, y + 18), name, fill="#172033", font=body_font)
            d.text((x + 58, y + 49), desc, fill="#536071", font=font(18))
            y += 115
    d.text((90, 810), "建议：生产环境尽量固定一种主路径。Docker 部署时重点确认端口映射、volume 位置和 .env 密钥；裸机部署时重点确认 systemd、数据库和 Redis。", fill="#26364d", font=body_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_oauth_sequence_image(path: Path):
    img = Image.new("RGB", (1600, 920), "#fbfcfe")
    d = ImageDraw.Draw(img)
    title_font = font(40, True)
    h_font = font(24, True)
    body_font = font(20)
    d.text((70, 45), "OpenAI/Codex OAuth 授权与刷新流程", fill="#172033", font=title_font)
    actors = [("管理员前端", 120), ("Sub2API 后端", 500), ("OpenAI 授权服务", 890), ("PostgreSQL", 1260)]
    for name, x in actors:
        rounded_box(d, (x, 145, x + 240, 205), "#eef4ff", "#4267d5", 18, 2)
        d.text((x + 38, 162), name, fill="#172033", font=h_font)
        d.line([(x + 120, 205), (x + 120, 820)], fill="#c5cfdf", width=3)
    steps = [
        (0, 1, "1. 请求生成授权链接"),
        (1, 2, "2. 拼接 OAuth URL"),
        (0, 2, "3. 浏览器跳转授权"),
        (2, 1, "4. 回调 code"),
        (1, 2, "5. code 换 token"),
        (1, 3, "6. 写入 Account.credentials JSON"),
        (1, 2, "7. 定时刷新 access token"),
        (1, 3, "8. 更新刷新结果/失败原因"),
    ]
    y = 255
    for a, b, text in steps:
        x1 = actors[a][1] + 120
        x2 = actors[b][1] + 120
        arrow(d, (x1, y), (x2, y), fill=(70, 88, 125), width=3)
        d.text((min(x1, x2) + 20, y - 28), text, fill="#26364d", font=body_font)
        y += 70
    rounded_box(d, (110, 820, 1490, 885), "#fff8e8", "#e3a943", 18, 2)
    d.text((140, 838), "准确性要点：OAuth 不是保存为某个 docker 文件，而是保存进数据库账号表的 credentials JSON；刷新模型列表和刷新 token 是两件不同的事。", fill="#26364d", font=font(19))
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_payment_flow_image(path: Path):
    img = Image.new("RGB", (1600, 900), "#f9fbfd")
    d = ImageDraw.Draw(img)
    title_font = font(40, True)
    h_font = font(24, True)
    body_font = font(20)
    d.text((70, 45), "支付订单与订阅履约流程", fill="#172033", font=title_font)
    steps = [
        ("选择套餐", "用户在 /purchase\n读取 plans/channels"),
        ("创建订单", "POST /payment/orders\n生成 out_trade_no"),
        ("跳转/扫码", "EasyPay/支付宝\n微信/Stripe"),
        ("支付回调", "/payment/webhook/*\n验签 + 幂等"),
        ("订单履约", "创建订阅/充值\n记录审计日志"),
        ("结果查询", "用户订单页\n管理员订单页"),
    ]
    x0, y0 = 70, 210
    for i, (title, body) in enumerate(steps):
        x = x0 + i * 250
        rounded_box(d, (x, y0, x + 220, y0 + 185), "#eaf9fb" if i % 2 else "#eef4ff", "#25a2b8" if i % 2 else "#4267d5", 22, 3)
        d.text((x + 24, y0 + 22), title, fill="#172033", font=h_font)
        wrap_draw_text(d, (x + 24, y0 + 72), body, 172, "#3d4a5f", body_font)
        if i < len(steps) - 1:
            arrow(d, (x + 222, y0 + 92), (x + 250 - 15, y0 + 92))
    risks = [
        ("幂等", "同一个回调可能重复到达，不能重复履约。"),
        ("验签", "必须确认回调确实来自支付平台。"),
        ("金额", "金额精度和币种不能用浮点随意计算。"),
        ("对账", "支付成功但履约失败要能后台重试。"),
    ]
    y = 520
    for i, (title, body) in enumerate(risks):
        x = 130 + i * 360
        rounded_box(d, (x, y, x + 300, y + 135), "#fff8e8", "#e3a943", 20, 2)
        d.text((x + 24, y + 22), title, fill="#172033", font=h_font)
        wrap_draw_text(d, (x + 24, y + 65), body, 250, "#536071", body_font)
    d.text((120, 760), "设计理解：支付模块把 provider 抽象出来，是为了以后新增支付通道时不改动订单主流程；订单状态机和回调幂等是安全重点。", fill="#26364d", font=body_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_scheduler_image(path: Path):
    img = Image.new("RGB", (1600, 880), "#fbfcff")
    d = ImageDraw.Draw(img)
    title_font = font(40, True)
    h_font = font(24, True)
    body_font = font(20)
    d.text((70, 45), "账号调度与 Failover 决策图", fill="#172033", font=title_font)
    nodes = [
        ("请求进入", 90, 180, "#eef4ff", "#4267d5", "API Key + model"),
        ("找到分组", 390, 180, "#e8f8ef", "#39a76d", "Group platform\nallowed models"),
        ("筛账号池", 690, 180, "#fff5e6", "#e49b32", "status/schedulable\nquota/rate limit"),
        ("选择账号", 990, 180, "#f1edff", "#7b61d1", "priority/load\nconcurrency"),
        ("转发上游", 1290, 180, "#eaf9fb", "#25a2b8", "stream/non-stream"),
        ("失败记录", 390, 540, "#ffecec", "#d44d4d", "timeout/429/5xx"),
        ("临时不可调度", 690, 540, "#fff8e8", "#e3a943", "overload_until\nrate_limit_reset"),
        ("重试/换账号", 990, 540, "#eef4ff", "#4267d5", "failover loop"),
    ]
    for title, x, y, fill, outline, body in nodes:
        rounded_box(d, (x, y, x + 220, y + 155), fill, outline, 22, 3)
        d.text((x + 24, y + 20), title, fill="#172033", font=h_font)
        wrap_draw_text(d, (x + 24, y + 66), body, 170, "#3d4a5f", body_font)
    for x in [310, 610, 910, 1210]:
        arrow(d, (x, 258), (x + 80, 258))
    arrow(d, (1400, 335), (500, 540), fill=(190, 80, 80))
    arrow(d, (610, 617), (690, 617), fill=(190, 80, 80))
    arrow(d, (910, 617), (990, 617), fill=(190, 80, 80))
    arrow(d, (1100, 540), (1100, 335), fill=(70, 88, 125))
    d.text((95, 780), "优化方向：后台展示“为什么这个 Key 没选到账号”的决策链，可以显著降低配置和排障成本。", fill="#26364d", font=body_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def parse_vue_routes() -> list[dict[str, str]]:
    text = read_text(ROOT / "frontend" / "src" / "router" / "index.ts")
    routes = []
    for block in re.finditer(r"\{\s*path:\s*'([^']+)'.*?\n\s*\}", text, re.S):
        b = block.group(0)
        path = re.search(r"path:\s*'([^']+)'", b)
        name = re.search(r"name:\s*'([^']+)'", b)
        comp = re.search(r"import\('([^']+)'\)", b)
        title = re.search(r"title:\s*'([^']+)'", b)
        requires_auth = "requiresAuth: false" not in b
        requires_admin = "requiresAdmin: true" in b
        requires_payment = "requiresPayment: true" in b
        if path:
            routes.append({
                "path": path.group(1),
                "name": name.group(1) if name else "",
                "component": comp.group(1) if comp else "",
                "title": title.group(1) if title else "",
                "auth": "否" if not requires_auth else "是",
                "admin": "是" if requires_admin else "否",
                "payment": "是" if requires_payment else "否",
            })
    seen = set()
    unique = []
    for r in routes:
        key = (r["path"], r["name"])
        if key not in seen:
            unique.append(r)
            seen.add(key)
    return unique


def parse_go_routes() -> list[dict[str, str]]:
    files = list((ROOT / "backend" / "internal" / "server" / "routes").glob("*.go"))
    result: list[dict[str, str]] = []
    for path in files:
        text = read_text(path)
        current_groups = {"r": "", "v1": "/api/v1"}
        for line in text.splitlines():
            stripped = line.strip()
            m_group = re.search(r"(\w+)\s*:=\s*(\w+)\.Group\(\"([^\"]*)\"\)", stripped)
            if m_group:
                var, parent, prefix = m_group.groups()
                current_groups[var] = (current_groups.get(parent, "") + "/" + prefix.strip("/")).replace("//", "/")
            m_call = re.search(r"(\w+)\.(GET|POST|PUT|DELETE|PATCH)\(\"([^\"]*)\",\s*([^,)]+)", stripped)
            if m_call:
                var, method, suffix, handler = m_call.groups()
                if var not in current_groups:
                    continue
                prefix = current_groups[var]
                result.append({
                    "file": str(path.relative_to(ROOT)).replace("\\", "/"),
                    "method": method,
                    "path": (prefix.rstrip("/") + "/" + suffix.lstrip("/")).replace("//", "/") or "/",
                    "handler": handler.strip(),
                })
    dedup = []
    seen = set()
    for item in result:
        key = (item["method"], item["path"], item["handler"])
        if key not in seen:
            dedup.append(item)
            seen.add(key)
    return dedup


def schema_summaries() -> list[tuple[str, str, str]]:
    mapping = {
        "User": ("用户表", "保存邮箱、密码哈希、角色、余额、并发限制、2FA、通知邮箱和登录活跃信息。"),
        "APIKey": ("API Key 表", "保存用户生成的 sk- 密钥、所属分组、状态、额度、IP 黑白名单和窗口用量。"),
        "Group": ("分组表", "定义模型平台、账号池、倍率、RPM 限制、订阅策略和用户可用范围。"),
        "Account": ("上游账号表", "保存 Claude/OpenAI/Gemini/Antigravity 等账号、凭据 JSON、代理、状态、过期和调度信息。"),
        "UsageLog": ("用量日志表", "记录每次请求的用户、Key、账号、模型、Token、成本、时延、图片数量和账单信息。"),
        "Proxy": ("代理表", "保存代理地址、类型、可用性、延迟、账号绑定和质量检测结果。"),
        "RedeemCode": ("兑换码表", "保存卡密额度、状态、过期时间和兑换记录。"),
        "SubscriptionPlan": ("订阅套餐表", "保存套餐名称、价格、周期、额度和可购买状态。"),
        "UserSubscription": ("用户订阅表", "保存用户订阅哪个分组、有效期、每日/周/月用量窗口和分配人。"),
        "PaymentOrder": ("支付订单表", "保存订单号、金额、支付方式、状态、履约和退款信息。"),
        "PaymentProviderInstance": ("支付通道表", "保存 EasyPay、支付宝、微信、Stripe 等支付配置。"),
        "Announcement": ("公告表", "保存公告标题、内容、发布状态、通知方式和可见范围。"),
        "ChannelMonitor": ("渠道监控表", "保存渠道健康检测任务、模型、Endpoint、频率和结果。"),
        "ChannelMonitorHistory": ("渠道监控历史表", "保存每次检测的成功失败、延迟和错误。"),
        "Setting": ("设置表", "保存系统设置、注册策略、支付开关、品牌展示、SMTP 等配置。"),
        "SecuritySecret": ("安全密钥表", "保存系统运行所需的安全密钥材料。"),
        "AuthIdentity": ("第三方登录身份表", "保存邮箱、LinuxDo、微信、OIDC 等外部身份和用户绑定关系。"),
        "PendingAuthSession": ("待完成登录会话表", "保存 OAuth 后还需要绑定邮箱/创建账号的临时流程。"),
        "PromoCode": ("优惠码表", "保存支付优惠码、折扣和使用限制。"),
        "ErrorPassthroughRule": ("错误透传规则表", "定义某些上游错误是否直接透传给客户端。"),
        "TLSFingerprintProfile": ("TLS 指纹模板表", "保存上游请求使用的 TLS/HTTP 指纹策略。"),
        "IdempotencyRecord": ("幂等记录表", "防止重复提交导致重复扣费、重复创建订单或重复操作。"),
    }
    summaries = []
    for file in sorted((ROOT / "backend" / "ent" / "schema").glob("*.go")):
        if file.name.endswith("_test.go"):
            continue
        text = read_text(file)
        m = re.search(r"type\s+(\w+)\s+struct", text)
        if not m:
            continue
        name = m.group(1)
        cn, desc = mapping.get(name, ("业务数据表", f"{name} 相关业务数据，字段定义见 ent/schema/{file.name}。"))
        fields = re.findall(r'field\.\w+\("([^"]+)"', text)
        summaries.append((name, cn, desc + " 主要字段：" + "、".join(fields[:12]) + ("..." if len(fields) > 12 else "")))
    return summaries


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, color="222222", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_layout(table, widths: list[float] | None):
    if not widths:
        table.autofit = True
        return

    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_twips = int(sum(widths) / 2.54 * 1440)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width_cm in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
        grid.append(col)
    tbl.insert(0, grid)


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    # The artifact DOCX renderer used in this environment can collapse wide
    # Word tables into vertical text. Use structured paragraphs instead: less
    # fancy than tables, but stable and readable in both Word and render QA.
    header = doc.add_paragraph()
    set_paragraph_shading(header, "274060")
    header.paragraph_format.left_indent = Cm(0.18)
    header.paragraph_format.right_indent = Cm(0.18)
    header.paragraph_format.space_before = Pt(4)
    header.paragraph_format.space_after = Pt(2)
    r = header.add_run("  /  ".join(headers))
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(255, 255, 255)

    for idx, row in enumerate(rows, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.28)
        p.paragraph_format.right_indent = Cm(0.15)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        if idx % 2 == 0:
            set_paragraph_shading(p, "F6F8FB")
        for i, value in enumerate(row):
            label = headers[i] if i < len(headers) else f"字段{i + 1}"
            rb = p.add_run(f"{label}：")
            rb.bold = True
            rb.font.name = "Microsoft YaHei"
            rb._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            rb.font.size = Pt(8.8 if len(rows) > 60 else 9.4)
            rb.font.color.rgb = RGBColor(35, 65, 105)
            rv = p.add_run(str(value))
            rv.font.name = "Microsoft YaHei"
            rv._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            rv.font.size = Pt(8.8 if len(rows) > 60 else 9.4)
            rv.font.color.rgb = RGBColor(34, 42, 55)
            if i != len(row) - 1:
                sep = p.add_run("    ")
                sep.font.size = Pt(8)
    doc.add_paragraph()
    return None


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if level == 1:
            run.font.color.rgb = RGBColor(30, 64, 105)
        elif level == 2:
            run.font.color.rgb = RGBColor(45, 83, 130)
    return p


def add_para(doc, text, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        rest = text[len(bold_prefix):]
        r2 = p.add_run(rest)
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for r in runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(34, 42, 55)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)


def add_code(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    for line in code.strip("\n").splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(37, 52, 78)


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("Sub2API 前后端功能与代码说明书")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(30, 64, 105)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("面向部署、运营、二次开发和非 TS/Go 背景阅读者")
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(83, 99, 124)
    doc.add_paragraph()
    meta = [
        ["项目路径", str(ROOT)],
        ["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["技术栈", "Vue 3 + TypeScript + Vite + Pinia；Go + Gin + Ent；PostgreSQL + Redis"],
        ["文档用途", "快速理解系统功能、前后端代码分工、核心数据流、部署与二次开发入口"],
    ]
    add_table(doc, ["项目", "说明"], meta, [3.2, 12])
    add_para(doc, "阅读建议：如果你不熟 TS 和 Go，先看“系统怎么跑起来”和“功能总览”，再看前端页面、后端接口、数据表，最后看代码解释章节。")
    doc.add_page_break()


def write_manual():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    arch = IMG_DIR / "architecture.png"
    flow = IMG_DIR / "gateway_flow.png"
    fmap = IMG_DIR / "frontend_map.png"
    domain_img = IMG_DIR / "domain_model.png"
    deploy_img = IMG_DIR / "deployment_topology.png"
    oauth_img = IMG_DIR / "oauth_sequence.png"
    payment_img = IMG_DIR / "payment_flow.png"
    scheduler_img = IMG_DIR / "scheduler_failover.png"
    make_architecture_image(arch)
    make_gateway_flow_image(flow)
    make_frontend_map_image(fmap)
    make_domain_model_image(domain_img)
    make_deployment_image(deploy_img)
    make_oauth_sequence_image(oauth_img)
    make_payment_flow_image(payment_img)
    make_scheduler_image(scheduler_img)

    vue_routes = parse_vue_routes()
    go_routes = parse_go_routes()
    schemas = schema_summaries()

    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    add_heading(doc, "图解阅读导航", 1)
    add_para(doc, "为了让这份文档不只是文字说明，本版增加了多张图：先看图可以快速建立整体模型，再回到对应章节看细节。")
    add_table(doc, ["图", "建议先理解的问题"], [
        ["总体架构图", "前端、后端、数据库、Redis、上游 AI、支付/备份之间如何协作。"],
        ["前端功能地图", "用户、管理员、支付、监控页面分别在哪里。"],
        ["一次模型请求流转图", "客户端请求如何经过鉴权、调度、协议转换、上游转发和计费。"],
        ["核心数据关系图", "User、APIKey、Group、Account、UsageLog、PaymentOrder 的关系。"],
        ["账号调度与 Failover 图", "为什么某个 API Key 或模型可能走不到账号。"],
        ["OAuth 时序图", "OpenAI/Codex OAuth 凭据为什么保存在数据库，不是文件。"],
        ["支付履约流程图", "订单从创建、回调、履约到查询的完整链路。"],
        ["部署形态对比图", "Docker、裸机 systemd、datamanagementd 三种部署组件如何区分。"],
    ], [4.2, 10.6])
    doc.add_page_break()

    add_heading(doc, "1. 项目一句话说明", 1)
    add_para(doc, "Sub2API 是一个 AI API 网关平台：用户拿平台生成的 API Key 调用模型，系统在后端完成鉴权、账号池调度、请求转发、用量统计、计费、监控和后台管理。")
    add_para(doc, "它不是单纯的网页项目，也不是单纯的后端接口项目，而是一个“管理后台 + API 中转网关 + 账号池调度 + 支付/订阅 + 运维监控”的完整平台。")
    add_bullets(doc, [
        "前端 frontend：Vue 页面，负责登录、用户中心、管理员后台、支付页面和可视化看板。",
        "后端 backend：Go 服务，负责路由、鉴权、业务逻辑、数据库访问、上游 AI 请求转发。",
        "数据库 PostgreSQL：保存用户、API Key、上游账号、订单、订阅、用量日志等核心业务数据。",
        "Redis：保存限流、缓存、并发、会话、调度状态等高频临时数据。",
        "deploy：保存 Docker Compose、裸机 systemd、安装脚本和示例配置。",
    ])
    doc.add_picture(str(arch), width=Inches(6.8))

    add_heading(doc, "2. 系统启动和代码入口", 1)
    add_table(doc, ["位置", "作用", "给非开发者的解释"], [
        ["backend/cmd/server/main.go", "Go 后端启动入口", "相当于后端程序的总开关，启动配置、数据库、Redis、路由和后台任务。"],
        ["backend/internal/server/router.go", "注册 HTTP 路由", "决定访问哪个 URL 时，后端由哪个处理函数响应。"],
        ["backend/internal/server/routes/*.go", "按功能拆分路由", "auth/user/admin/payment/gateway 分别管理不同业务入口。"],
        ["frontend/src/main.ts", "Vue 前端启动入口", "把 App.vue、路由、状态管理挂到浏览器页面上。"],
        ["frontend/src/router/index.ts", "前端页面路由", "决定 /login、/admin/accounts 等地址显示哪个页面组件。"],
        ["frontend/src/api/client.ts", "前端请求封装", "所有前端调用后端接口，基本都会经过这里统一加 token、处理错误。"],
        ["frontend/src/stores/*.ts", "前端状态仓库", "保存登录状态、系统设置、支付状态、公告等跨页面共享数据。"],
    ], [4.2, 4.4, 7.2])

    add_heading(doc, "3. 前端功能总览", 1)
    add_para(doc, "前端采用 Vue 3。你可以把 Vue 页面理解成一个个“可组合的页面积木”：views 是页面，components 是复用组件，api 是请求后端的函数，stores 是全局状态。")
    doc.add_picture(str(fmap), width=Inches(6.8))
    frontend_modules = [
        ["公开/初始化", "SetupWizard、Home、Login、Register、EmailVerify、OAuth callbacks、Forgot/Reset password", "首次部署初始化、登录注册、第三方登录回调、找回密码。"],
        ["用户中心", "Dashboard、Keys、Usage、Redeem、Profile、Subscriptions、AvailableChannels", "普通用户查看余额、生成 API Key、看调用记录、兑换卡密、看套餐和个人信息。"],
        ["管理员后台", "Dashboard、Users、Groups、Accounts、Channels、Settings、Ops、Backup", "管理员维护用户、分组、上游账号、模型渠道、系统设置、备份和运维监控。"],
        ["支付订单", "PaymentView、PaymentQRCode、StripePayment、UserOrders、AdminOrders、PaymentPlans", "用户买套餐/充值，管理员配置支付通道、套餐和订单处理。"],
        ["监控扩展", "ChannelStatus、ChannelMonitor、CustomPage、Announcements", "渠道健康检测、公告通知、自定义菜单页面。"],
    ]
    add_table(doc, ["功能域", "主要页面/组件", "说明"], frontend_modules, [3, 6.3, 6.2])

    add_heading(doc, "4. 前端路由清单", 1)
    add_para(doc, "下面来自 frontend/src/router/index.ts。auth=是否需要登录，admin=是否必须管理员，payment=是否依赖支付功能开关。")
    route_rows = [[r["path"], r["name"], r["component"].replace("@/", "src/"), r["auth"], r["admin"], r["payment"]] for r in vue_routes]
    add_table(doc, ["路径", "路由名", "页面组件", "登录", "管理员", "支付"], route_rows, [2.4, 3.0, 5.4, 1.1, 1.2, 1.1])

    add_heading(doc, "5. 前端代码怎么读", 1)
    add_para(doc, "TypeScript 是 JavaScript 的加强版，主要多了“类型”。类型不是业务逻辑，它像标签，告诉编辑器某个变量应该是什么形状，减少写错字段名、传错参数的问题。")
    add_code(doc, """
// 一个典型前端 API 函数长这样：
export function listUsers(params: QueryParams) {
  return apiClient.get('/admin/users', { params })
}

// 用人话翻译：
// 1. listUsers 是“查询用户列表”函数
// 2. params 是查询条件，比如页码、关键字
// 3. apiClient.get 表示向后端 GET /admin/users 发请求
""")
    add_para(doc, "Vue 单文件组件 .vue 通常由三部分组成：template 是页面结构，script 是数据和方法，style 是样式。项目里很多页面会在 script 里调用 src/api 下的函数，然后把结果显示到 template。")
    add_code(doc, """
<template>
  <div>{{ user.email }}</div>
</template>

<script setup lang="ts">
// script setup 是 Vue 3 的简洁写法
// lang=\"ts\" 表示这里用 TypeScript
</script>
""")

    add_heading(doc, "6. 后端功能总览", 1)
    add_para(doc, "后端采用 Go。你可以把后端分成五层：routes 接 URL，handler 处理 HTTP 输入输出，service 写业务规则，repository 读写数据库/Redis，ent/schema 定义数据库表。")
    doc.add_picture(str(flow), width=Inches(6.8))
    backend_layers = [
        ["routes", "backend/internal/server/routes", "注册 URL 和 HTTP 方法，例如 POST /auth/login。"],
        ["middleware", "backend/internal/server/middleware", "在进入业务前做鉴权、CORS、安全头、限流、请求体大小、日志等检查。"],
        ["handler", "backend/internal/handler", "把 HTTP 请求转成业务参数，调用 service，再把结果转成 JSON。"],
        ["service", "backend/internal/service", "最核心的业务逻辑：账号调度、计费、订阅、支付、OAuth、监控、网关转发。"],
        ["repository", "backend/internal/repository", "数据库、Redis、外部 HTTP 客户端、GitHub Release、S3 等基础访问。"],
        ["ent/schema", "backend/ent/schema", "数据库表结构定义。修改这里后需要重新生成 Ent 代码。"],
    ]
    add_table(doc, ["层", "代码位置", "作用"], backend_layers, [2.5, 5.2, 7.5])

    add_heading(doc, "7. 后端路由与接口分组", 1)
    route_groups = [
        ["认证 auth", "/api/v1/auth/*", "注册、登录、刷新 token、退出、验证码、OAuth 登录/绑定、找回密码。"],
        ["用户 user", "/api/v1/user、/keys、/usage、/redeem、/subscriptions", "普通用户资料、API Key、用量、兑换码、订阅、公告。"],
        ["管理员 admin", "/api/v1/admin/*", "用户、分组、账号、公告、代理、设置、备份、监控、系统升级、渠道等后台管理。"],
        ["支付 payment", "/api/v1/payment/*", "用户下单、查订单、取消、退款申请、支付回调、管理员订单/套餐/通道管理。"],
        ["AI 网关 gateway", "/v1、/v1beta、/responses、/chat/completions、/images/*", "兼容 Claude/OpenAI/Gemini/Codex/Antigravity 客户端请求。"],
        ["通用 common", "/health、/setup/status", "健康检查和部署初始化状态。"],
    ]
    add_table(doc, ["接口组", "典型路径", "功能"], route_groups, [3.2, 5.0, 7.0])
    add_para(doc, "接口清单较长，下面列出从 routes/*.go 抽取到的主要路由。注意：部分复杂路由通过函数闭包或嵌套组装，表格用于理解入口，不代替正式 OpenAPI 文档。")
    go_rows = [[r["method"], r["path"], r["handler"], r["file"]] for r in go_routes[:180]]
    add_table(doc, ["方法", "路径", "处理函数", "来源文件"], go_rows, [1.3, 5.0, 4.6, 4.8])

    add_heading(doc, "8. AI 网关核心能力", 1)
    gateway_features = [
        ["多协议兼容", "支持 Claude Messages、OpenAI Chat Completions、OpenAI Responses、Images、Gemini v1beta、Antigravity 专用入口。"],
        ["API Key 鉴权", "客户端不直接拿上游账号 token，而是使用平台生成的 sk- key。后端验证 key、用户、分组、订阅和额度。"],
        ["账号池调度", "按平台、分组、模型、账号状态、并发、优先级、限流、临时不可调度等条件选择上游账号。"],
        ["模型映射", "允许把客户端请求模型映射到上游真实模型，尤其是 OpenAI/Codex、Antigravity、Gemini 场景。"],
        ["Failover", "上游失败时尝试切换账号或记录不可用原因，避免单个账号故障拖垮服务。"],
        ["计费与日志", "统计输入/输出/cache token、图片数量、耗时、成本、倍率，写入 usage_logs。"],
        ["运维观测", "Ops 模块展示 QPS、延迟、错误、上游失败、系统日志、OpenAI token 统计等。"],
    ]
    add_table(doc, ["能力", "说明"], gateway_features, [3.2, 12.2])
    add_para(doc, "结合你前面问过的 OpenAI OAuth：OpenAI/Codex 的 OAuth 凭据最终进入账号 credentials JSON 字段，不是一个单独文件。请求时由 token provider/refresh service 从账号凭据里取 token、刷新 token、再转发给上游。")
    doc.add_picture(str(scheduler_img), width=Inches(6.8))
    doc.add_picture(str(oauth_img), width=Inches(6.8))

    add_heading(doc, "9. 用户、分组、账号、Key 的关系", 1)
    add_para(doc, "这是理解 Sub2API 最关键的一组关系：用户不是直接调用上游账号，而是先创建 API Key；API Key 绑定分组；分组绑定一批上游账号；请求进来后，后端从这个分组的账号池里挑一个可用账号。")
    doc.add_picture(str(domain_img), width=Inches(6.8))
    add_table(doc, ["对象", "通俗解释", "代码/表"], [
        ["User", "平台登录用户，有余额、角色、状态、并发限制。", "ent/schema/user.go"],
        ["APIKey", "用户真正给客户端使用的 sk- 密钥。", "ent/schema/api_key.go"],
        ["Group", "一组模型/账号池/倍率/限制策略。", "ent/schema/group.go"],
        ["Account", "真实上游账号，例如 OpenAI OAuth、Claude OAuth、Gemini OAuth 或 API Key。", "ent/schema/account.go"],
        ["AccountGroup", "账号和分组的关联表，一个账号可进多个分组。", "ent/schema/account_group.go"],
        ["UsageLog", "每次调用产生的用量和成本记录。", "ent/schema/usage_log.go"],
    ], [2.6, 8.3, 4.0])
    add_para(doc, "示例流程：用户在前端创建 Key -> 选择可用分组 -> 客户端拿 Key 请求 /v1/messages -> 后端查 Key 所属分组 -> 找可用 Account -> 转发上游 -> 记录 UsageLog -> 扣减额度或计入订阅用量。")

    add_heading(doc, "10. 数据表说明", 1)
    schema_rows = [[name, cn, desc] for name, cn, desc in schemas]
    add_table(doc, ["Schema", "中文含义", "主要用途"], schema_rows, [3.3, 3.5, 8.2])

    add_heading(doc, "11. 管理员后台功能详解", 1)
    admin_features = [
        ["仪表盘", "展示整体用量、趋势、模型统计、分组统计、用户排行、API Key 趋势和用户维度用量。"],
        ["用户管理", "创建、编辑、删除用户；调整余额；查看用户 Key、用量、订阅、属性和 RPM 状态。"],
        ["分组管理", "定义不同平台/模型池；设置倍率、RPM 覆盖、容量摘要、账号排序和可用范围。"],
        ["账号管理", "新增上游账号；OAuth 授权；刷新 token；批量创建/导入/导出；测试账号；清错误；刷新 tier；设置调度状态。"],
        ["渠道与监控", "配置模型渠道、默认价格、渠道健康检测、检测模板、手动运行和历史结果。"],
        ["代理管理", "维护代理服务器，测试连通性，质量检查，绑定账号和批量导入。"],
        ["支付管理", "配置支付开关、支付通道、套餐、订单、退款、履约重试和支付统计。"],
        ["公告管理", "创建面向用户的公告，支持发布状态、可见范围、已读统计。"],
        ["卡密/优惠码", "生成兑换码、批量删除、导出、过期；维护支付优惠码和使用记录。"],
        ["系统设置", "站点品牌、注册策略、SMTP、Turnstile、支付、OAuth、CORS、安全头、模型策略等。"],
        ["备份与数据管理", "S3 备份配置、定时备份、恢复；datamanagementd agent 健康检查和备份任务。"],
        ["运维 Ops", "实时并发、QPS、延迟、错误日志、告警规则、系统日志、运行时日志配置。"],
    ]
    add_table(doc, ["后台模块", "功能说明"], admin_features, [3.2, 12.2])

    add_heading(doc, "12. 普通用户功能详解", 1)
    user_features = [
        ["用户仪表盘", "查看自己的余额、用量趋势、模型分布、API Key 使用情况。"],
        ["API Key 管理", "创建、编辑、删除 Key；选择分组；配置额度和 IP 黑白名单。"],
        ["用量记录", "查看每次调用消耗的模型、Token、费用、时间和关联 Key。"],
        ["兑换码", "输入管理员发放的卡密，给账户充值或获得权益。"],
        ["订阅", "查看当前订阅、额度进度、过期时间和可用分组。"],
        ["支付购买", "选择套餐，创建订单，扫码/Stripe 支付，查看订单结果。"],
        ["个人资料", "修改密码、绑定邮箱/第三方身份、配置 TOTP 双因素认证和通知邮箱。"],
        ["可用渠道", "查看自己当前可访问的模型渠道和状态。"],
    ]
    add_table(doc, ["用户模块", "功能说明"], user_features, [3.2, 12.2])

    add_heading(doc, "13. 支付和订阅逻辑", 1)
    add_para(doc, "支付模块内置在项目里，不需要单独部署 Sub2ApiPay。它支持用户创建订单、选择套餐、支付回调、订单履约、退款申请和管理员处理。")
    doc.add_picture(str(payment_img), width=Inches(6.8))
    add_bullets(doc, [
        "用户侧接口在 /api/v1/payment：获取套餐、获取支付通道、创建订单、验证订单、取消订单、申请退款。",
        "公开接口在 /api/v1/payment/public：用于支付完成后的无登录恢复和查询。",
        "回调接口在 /api/v1/payment/webhook：EasyPay、支付宝、微信、Stripe 会调用这里通知支付结果。",
        "管理员接口在 /api/v1/admin/payment：配置支付、管理订单、套餐和支付通道。",
        "支付成功后通常会触发 fulfillment，将套餐转换为用户订阅、余额或权益。",
    ])

    add_heading(doc, "14. 认证、安全和权限", 1)
    add_table(doc, ["机制", "说明"], [
        ["JWT 登录态", "用户登录后前端保存 token，访问受保护页面时带上 token。"],
        ["Refresh Token", "用于刷新登录态，后端支持登出和撤销所有会话。"],
        ["管理员中间件", "访问 /admin/* 必须是管理员角色。"],
        ["API Key 中间件", "访问 /v1、/v1beta、/responses 等网关接口必须带有效 sk- key。"],
        ["Rate Limit", "注册、登录、验证码、找回密码、OAuth 待完成流程有 Redis 限流，Redis 故障时偏 fail-close。"],
        ["TOTP 2FA", "用户可启用双因素认证，密钥加密后保存。"],
        ["CORS/安全头", "生产环境建议配置 allowed_origins、trusted_proxies、CSP、Turnstile 等。"],
        ["URL Allowlist", "对外部 URL、上游地址、价格源等做白名单/HTTP 限制，降低 SSRF 和明文传输风险。"],
    ], [3.5, 11.5])

    add_heading(doc, "15. Go 代码怎么读", 1)
    add_para(doc, "Go 文件以 package 开头，同一个目录通常属于同一个包。函数名首字母大写表示可被其他包调用，小写表示包内使用。")
    add_code(doc, """
func RegisterUserRoutes(v1 *gin.RouterGroup, h *handler.Handlers, jwtAuth middleware.JWTAuthMiddleware, settingService *service.SettingService) {
    authenticated := v1.Group(\"\")
    authenticated.Use(gin.HandlerFunc(jwtAuth))
    user := authenticated.Group(\"/user\")
    user.GET(\"/profile\", h.User.GetProfile)
}

// 用人话翻译：
// 1. RegisterUserRoutes 是“注册用户相关接口”
// 2. authenticated.Use(jwtAuth) 表示这些接口必须登录
// 3. GET /user/profile 会交给 h.User.GetProfile 处理
""")
    add_para(doc, "handler 负责 HTTP，service 负责业务。比如登录接口：handler 读取邮箱密码 -> 调 AuthService -> service 查用户/校验密码/生成 token -> handler 返回 JSON。")

    add_heading(doc, "16. Ent 和数据库怎么理解", 1)
    add_para(doc, "Ent 是 Go 的 ORM。ORM 可以理解为“用 Go 代码描述数据库表”，然后自动生成查询、创建、更新、删除代码。backend/ent/schema 里的每个文件大多对应一张表。")
    add_code(doc, """
field.String(\"email\").Unique()

// 用人话翻译：
// 数据库 users 表里有一个 email 字段，
// 类型是字符串，并且不能重复。
""")
    add_para(doc, "如果改了 ent/schema 下的字段，不能只保存文件，还要执行 go generate ./ent 重新生成 ent 代码，否则后端可能编译不过或数据库结构不一致。")

    add_heading(doc, "17. 部署相关说明", 1)
    doc.add_picture(str(deploy_img), width=Inches(6.8))
    add_table(doc, ["部署方式", "文件", "说明"], [
        ["Docker Compose", "deploy/docker-compose.yml", "启动 sub2api、PostgreSQL、Redis，适合服务器部署。"],
        ["本地目录 Compose", "deploy/docker-compose.local.yml（当前仓库可能未包含）", "把 postgres_data、redis_data、data 直接放部署目录，方便备份迁移。"],
        ["裸机 systemd", "deploy/install.sh、deploy/sub2api.service", "把二进制放 /opt/sub2api，用 systemd 启动。"],
        ["数据管理 agent", "deploy/install-datamanagementd.sh、sub2api-datamanagementd.service", "宿主机运行 datamanagementd，提供备份/数据管理 socket。"],
        ["镜像构建", "deploy/build_image.sh、Dockerfile", "本地构建 Docker 镜像。"],
    ], [3.2, 5.2, 7.0])
    add_para(doc, "结合你前面部署时遇到的问题：Docker 端口要保证容器内应用监听端口和 compose 映射一致，例如容器内 SERVER_PORT=8080 时，ports 可写 18080:8080；如果应用实际监听 18080，则应映射 18080:18080。")

    add_heading(doc, "18. 二次开发常见入口", 1)
    add_table(doc, ["你想改什么", "优先看哪里", "注意点"], [
        ["新增一个前端页面", "frontend/src/views + frontend/src/router/index.ts", "加页面组件，再加路由；如需菜单，还要看导航组件/设置。"],
        ["新增前端调用接口", "frontend/src/api/*.ts", "统一走 apiClient，保持 token 和错误处理一致。"],
        ["新增后端接口", "backend/internal/server/routes + handler + service", "先注册路由，再写 handler，复杂业务放 service。"],
        ["新增数据库字段", "backend/ent/schema/*.go", "改 schema 后执行 go generate ./ent，并考虑迁移。"],
        ["新增上游平台", "gateway/service/account/token provider/model mapping", "这是高风险改动，需要处理鉴权、模型、计费、错误、流式响应。"],
        ["新增支付通道", "payment_* service/handler/repository", "要处理签名校验、回调幂等、订单状态机和退款。"],
        ["改部署配置", "deploy/*.yml、Dockerfile、config.example.yaml", "注意端口、数据卷、密钥和数据库连接。"],
    ], [3.3, 5.4, 6.4])

    add_heading(doc, "19. 重要风险和维护建议", 1)
    add_bullets(doc, [
        "账号凭据、JWT_SECRET、TOTP_ENCRYPTION_KEY、支付密钥必须妥善保存，不能提交到公开仓库。",
        "PostgreSQL 是核心数据来源，Docker named volume 或本地 postgres_data 必须定期备份。",
        "Redis 不是永久数据库，但限流、缓存和调度依赖它，生产环境 Redis 不稳定会影响登录、网关和后台体验。",
        "上游模型名更新可能快于代码内置映射，需要通过后台模型映射或代码更新兜底。",
        "生产环境必须配置 CORS allowed_origins、trusted_proxies 和反向代理安全策略。",
        "高并发时关注 Ops 看板里的 QPS、延迟、错误分布、账号可用性和数据库写入压力。",
        "前端构建在低内存服务器上可能 OOM，建议开启 swap 或提高构建机内存，并避免生产服务器无 swap。"
    ])

    add_heading(doc, "20. 快速查找索引", 1)
    add_table(doc, ["问题", "看哪里"], [
        ["OpenAI OAuth token 存哪里", "backend/ent/schema/account.go 的 credentials 字段；后台 OpenAI OAuth handler/service。"],
        ["为什么请求走不到模型", "API Key -> Group -> Account 绑定关系；账号状态；模型映射；Ops 错误日志。"],
        ["数据库数据在哪", "Docker compose volume 或 postgres_data；裸机是 PostgreSQL 自己的数据目录。"],
        ["前端页面在哪改", "frontend/src/views，对应路由在 frontend/src/router/index.ts。"],
        ["接口在哪注册", "backend/internal/server/routes/*.go。"],
        ["业务逻辑在哪", "backend/internal/service/*.go。"],
        ["数据库表在哪定义", "backend/ent/schema/*.go。"],
        ["部署配置在哪", "deploy/docker-compose.yml、Dockerfile、deploy/.env。"],
    ], [5.0, 10.0])

    add_heading(doc, "21. 本次文档自查与遗漏补充", 1)
    add_para(doc, "我按前端目录、后端目录、路由文件、Ent 数据表和支撑模块重新核对了多轮。上一版已经覆盖主流程，但对“支撑型模块”“路由父路径”“图解表达”和“设计取舍”的解释不够细，本版补充如下。")
    add_table(doc, ["审校轮次", "检查重点", "发现/处理"], [
        ["第 1 轮", "目录覆盖", "确认 frontend/src/views、api、stores 与 backend/internal/handler、service、repository、ent/schema 已覆盖主功能。"],
        ["第 2 轮", "遗漏模块", "补充 frontend/components、composables、utils、i18n、types，以及 backend/internal/pkg、payment、setup、web、util。"],
        ["第 3 轮", "路由准确性", "修正静态扫描逻辑，让 admin/user/payment 等子路由继承 /api/v1 父路径；复杂闭包路由保留来源文件提示。"],
        ["第 4 轮", "设计取舍", "新增技术选型、为什么这样设计、优缺点、优化路线，避免只罗列功能。"],
        ["第 5 轮", "图解表达", "新增核心数据关系、OAuth、支付、部署、账号调度/Failover 等图，减少纯文字密度。"],
        ["第 6 轮", "版式渲染", "保留结构化卡片替代复杂 Word 表格，避免渲染器把宽表压成竖排；每次生成后渲染 PNG 检查。"],
    ], [2.6, 4.4, 8.0])
    add_table(doc, ["补充点", "代码位置", "为什么重要", "本版处理"], [
        ["前端复用组件", "frontend/src/components", "大量页面不是直接写完整 UI，而是复用 DataTable、Dialog、StatusBadge、支付组件、账号组件、图表组件。理解组件层有助于知道页面为什么看起来统一。", "新增组件层设计说明、优缺点和优化建议。"],
        ["前端组合函数", "frontend/src/composables", "OAuth、自动刷新、表格加载、分页大小、剪贴板、路由预加载等逻辑被抽成 composable，页面会更薄。", "新增组合函数设计原因与维护建议。"],
        ["前端工具函数", "frontend/src/utils", "错误格式化、价格展示、平台颜色、Key 脱敏、设备判断等都在这里，不全在页面里。", "新增 utils/常量/types/i18n 的职责说明。"],
        ["后端 internal/pkg", "backend/internal/pkg", "这里包含 OpenAI/Gemini/Antigravity 协议模型、OAuth、错误封装、Web Search、TLS 指纹、日志等底层能力。", "新增“协议兼容与上游适配层”说明。"],
        ["后端 internal/payment", "backend/internal/payment", "支付不是只在 service 里写死，而是抽象了 provider、registry、load_balancer、fee、amount 等通用能力。", "新增支付通道技术设计说明。"],
        ["后端 setup/web/util", "backend/internal/setup、web、util", "首次安装向导、嵌入前端静态资源、URL 校验、日志脱敏等属于平台完整性能力。", "新增部署和安全支撑模块说明。"],
        ["测试覆盖线索", "frontend/src/**/__tests__、backend/**/*_test.go", "项目有较多前后端测试文件，说明很多复杂逻辑已有单元/集成验证。", "新增测试与质量保障章节。"],
    ], [3.0, 4.2, 5.4, 3.4])

    add_heading(doc, "22. 技术选型详解", 1)
    add_para(doc, "这个项目的技术栈不是随便堆起来的，它服务于两个目标：一是管理后台要开发快、交互复杂但可维护；二是 API 网关要稳定、并发能力强、容易部署成单个二进制或容器。")
    tech_choices = [
        ["Vue 3", "适合后台管理系统，组合式 API 方便拆复用逻辑，生态成熟。", "学习成本比纯 HTML 高；大型页面如果不拆组件会变复杂。", "继续把复杂页面拆成更小组件，减少单文件超过千行的情况。"],
        ["TypeScript", "给前端请求、响应、表单、设置项加类型保护，减少字段写错。", "初学者会觉得类型定义啰嗦，接口变动时类型也要同步改。", "从后端生成 OpenAPI/类型，减少手写类型漂移。"],
        ["Vite", "开发启动快，构建现代前端项目体验好。", "构建大型项目时仍然吃内存，低配服务器上可能 OOM。", "生产构建放到 CI 或高内存机器，Docker 构建时控制 sourcemap/checker。"],
        ["Pinia", "比传统 Vuex 更轻，适合保存登录态、设置、支付状态、公告等全局数据。", "状态过多时可能变成隐形全局依赖。", "为 store 增加职责边界和持久化策略说明。"],
        ["Axios", "统一封装 token、错误处理、刷新登录态，前端调用后端更一致。", "如果每个 API 文件各自处理错误，会出现重复逻辑。", "保持 api/client.ts 为唯一请求入口，并统一错误码映射。"],
        ["Go", "编译成单个二进制，部署简单，性能和并发模型适合网关。", "泛型和错误处理对新手不如脚本语言直观。", "保持 handler/service/repository 分层，避免业务逻辑散落。"],
        ["Gin", "轻量、性能好，路由中间件清晰，适合 API 服务。", "没有像大型框架那样强约束目录结构，团队需要自律。", "继续保持 routes 按功能拆分，补充 OpenAPI 文档。"],
        ["Ent", "用 Go 代码定义表结构，生成类型安全的数据库访问代码。", "生成代码多，初学者容易不知道哪些文件可改、哪些是生成物。", "文档中明确：优先改 ent/schema，生成文件不要手工改。"],
        ["PostgreSQL", "适合保存强一致业务数据，支持 JSONB、索引和复杂查询。", "需要备份、迁移、连接池和慢查询治理。", "增加迁移规范、备份恢复演练和关键表索引审计。"],
        ["Redis", "适合缓存、限流、并发计数、调度状态等高频临时数据。", "Redis 故障会影响限流/登录/调度，不能当永久数据源。", "生产部署建议持久化、监控内存、设置合理淘汰策略。"],
        ["Docker Compose", "把后端、数据库、Redis 一起编排，用户部署成本低。", "端口、volume、env 配错时排查成本高。", "提供 production/local 两套清晰 compose 和数据目录说明。"],
    ]
    add_table(doc, ["技术", "为什么选", "缺点/代价", "优化建议"], tech_choices, [2.4, 5.0, 4.2, 4.2])

    add_heading(doc, "23. 核心设计为什么这样做", 1)
    design_rows = [
        ["前后端分离", "前端专注交互，后端专注 API 和网关，开发边界清楚。", "页面可以独立迭代，后端可同时服务网页和 API 客户端。", "部署和接口版本管理更复杂。", "补充 OpenAPI/接口契约测试，避免前后端字段漂移。"],
        ["handler/service/repository 分层", "HTTP、业务、数据访问职责分开，复杂业务不堆在路由里。", "利于测试和维护，换数据库或外部服务时影响面小。", "文件数量变多，新人要先理解调用链。", "每个核心 service 增加 README 或包注释。"],
        ["用户 -> API Key -> 分组 -> 账号池", "把客户身份、权限策略和上游账号解耦。", "一个用户可有多个 Key，一个 Key 可绑定不同分组，账号池可灵活调度。", "关系链长，配置错误时用户会觉得“模型不可用”。", "后台增加配置诊断：某 Key 为什么不可调度到模型。"],
        ["账号凭据存在 Account.credentials JSON", "不同平台凭据结构不同，用 JSONB 能兼容 OAuth/API Key/Cookie 等多形态。", "扩展平台方便，不必每加平台就加很多列。", "JSON 字段不如强类型列直观，查询和校验更依赖代码。", "为各平台 credentials 增加 schema 校验和脱敏展示。"],
        ["Redis 做限流和缓存", "限流计数、token 缓存、并发状态需要高频读写，数据库不适合承受全部压力。", "性能好，降低数据库压力。", "Redis 不稳定会影响高风险入口，fail-close 可能导致用户暂时无法登录。", "关键限流策略增加降级开关和监控告警。"],
        ["OAuth 账号自动刷新", "OpenAI/Gemini/Antigravity 等 OAuth token 会过期，必须后台刷新。", "减少人工维护，账号更稳定。", "刷新失败原因复杂，和上游策略强相关。", "后台显示 token 到期时间、刷新失败原因和一键重授权建议。"],
        ["Ops 运维看板", "网关业务最怕不知道哪里慢、哪里错、哪个账号坏。", "可视化 QPS、错误、延迟、账号状态，便于生产排障。", "统计本身会增加写入和查询压力。", "冷热数据分层，长期指标做预聚合，减少大表实时扫。"],
        ["内置支付和订阅", "平台希望自助充值、购买套餐、自动履约，不依赖独立支付系统。", "商业化闭环完整，用户自助体验好。", "支付安全、回调幂等、退款和对账会增加复杂度。", "增加对账任务、订单异常队列和支付通道健康检测。"],
        ["datamanagementd 独立 agent", "容器内直接操作宿主机备份/恢复能力有限，用宿主机 agent 处理更合适。", "权限边界更清楚，Docker 场景也能做本地备份。", "需要额外安装和挂载 socket，部署多一步。", "compose 脚本自动检测 agent/socket 并给出提示。"],
    ]
    add_table(doc, ["设计点", "为什么这样设计", "优点", "缺点/代价", "可优化点"], design_rows, [2.7, 4.0, 3.5, 3.5, 3.8])

    add_heading(doc, "24. 模块级优缺点和优化路线", 1)
    module_rows = [
        ["认证与用户", "覆盖邮箱登录、验证码、OAuth、OIDC、微信、LinuxDo、TOTP、会话撤销。", "功能完整，适合公开注册和私有部署。", "流程分支多，登录/绑定/创建账号容易出现边界状态。", "画出 OAuth 状态机；给每个登录来源增加端到端测试。"],
        ["API Key 管理", "支持额度、状态、分组、IP 黑白名单、使用统计。", "权限和计费粒度细。", "普通用户可能不理解 Key、分组、模型的关系。", "新增“使用向导”和“Key 可用性诊断”。"],
        ["账号池调度", "考虑平台、模型、状态、并发、限流、优先级、临时不可调度。", "能提高资源利用率和稳定性。", "调度规则复杂，线上问题不易复现。", "记录每次调度决策链路，后台可查看为什么选/不选某账号。"],
        ["协议兼容层", "兼容 Claude、OpenAI、Responses、Gemini、Antigravity 等不同协议。", "客户端接入灵活，用户迁移成本低。", "协议变化快，映射代码维护成本高。", "建立协议兼容测试集，跟随上游模型/字段变更。"],
        ["计费与用量", "记录 token、cache、图片、成本、倍率、订阅消耗。", "账单可追踪，适合商业化。", "流式请求、失败重试、缓存 token 的计费边界容易复杂。", "为每种 request_type 建立计费说明和回归测试。"],
        ["支付系统", "provider 抽象支持多支付方式，订单和履约分层。", "扩展支付通道较方便。", "第三方回调差异大，签名和时区金额处理容易踩坑。", "加强金额精度、回调重放、对账和退款状态机测试。"],
        ["前端后台页面", "按用户/管理员/支付/运维拆页面，配合通用组件。", "功能入口清楚，组件复用较多。", "部分页面功能重，维护时需要先理解很多弹窗和表格组件。", "继续拆分大页面；为复杂页面写“页面数据流图”。"],
        ["数据模型", "Ent schema 覆盖用户、Key、账号、订单、订阅、监控、设置等。", "模型比较完整，适合平台型业务。", "表多且关系复杂，直接看生成代码会迷路。", "维护 ER 图和关键表生命周期说明。"],
        ["部署体系", "Docker、裸机 systemd、data agent 都有支持。", "适配不同用户部署习惯。", "部署路径多，文档不统一时容易混淆。", "明确推荐路径：生产优先 docker-compose.local，本地开发走源码。"],
    ]
    add_table(doc, ["模块", "当前能力", "优点", "短板", "优化路线"], module_rows, [2.7, 4.1, 3.3, 3.5, 4.0])

    add_heading(doc, "25. 支撑目录补充说明", 1)
    support_rows = [
        ["frontend/src/components", "通用 UI、账号弹窗、支付组件、图表、布局、用户资料、监控卡片。", "页面不要重复造轮子，统一交互和样式。"],
        ["frontend/src/composables", "useOpenAIOAuth、useTableLoader、useRoutePrefetch、useClipboard 等复用逻辑。", "让页面代码更短，也方便测试单独逻辑。"],
        ["frontend/src/utils", "错误处理、格式化、价格、平台颜色、Key 脱敏、设备判断、功能开关。", "把纯函数从页面抽走，减少重复。"],
        ["frontend/src/i18n", "多语言文案。", "页面标题和按钮文案可以国际化。"],
        ["frontend/src/types", "前端业务类型。", "帮助 TypeScript 检查接口字段。"],
        ["backend/internal/pkg/apicompat", "Anthropic、OpenAI Responses、Chat Completions 之间的协议转换。", "这是多客户端兼容的关键。"],
        ["backend/internal/pkg/openai/gemini/antigravity", "上游平台的常量、请求结构、OAuth、模型列表和客户端封装。", "把平台差异隔离在 pkg 层。"],
        ["backend/internal/pkg/tlsfingerprint/proxyurl/urlvalidator", "TLS 指纹、代理解析、URL 安全校验。", "服务对外请求时降低封号、SSRF、错误代理格式风险。"],
        ["backend/internal/payment", "支付 provider、registry、amount、fee、load_balancer。", "支付通道可扩展，金额和手续费逻辑集中。"],
        ["backend/internal/setup", "首次安装向导和初始化配置。", "部署后能通过网页完成数据库/Redis/管理员初始化。"],
        ["backend/internal/web", "嵌入或提供前端静态资源。", "Go 二进制可以直接带上前端页面。"],
        ["backend/internal/util", "日志脱敏、响应头、URL 校验等工具。", "安全和通用逻辑集中维护。"],
    ]
    add_table(doc, ["目录", "里面有什么", "为什么要单独存在"], support_rows, [4.2, 6.2, 5.0])

    add_heading(doc, "26. 测试和质量保障", 1)
    add_para(doc, "项目里前端有 Vitest 测试，后端有大量 Go 单元测试和集成测试。对这种网关平台来说，测试的价值很高，因为很多 bug 不会在页面上立刻暴露，而是在某个上游协议、计费边界、OAuth 刷新或支付回调里出现。")
    add_table(doc, ["测试对象", "代码线索", "建议重点"], [
        ["前端组件测试", "frontend/src/**/__tests__/*.spec.ts", "登录弹窗、支付状态、表格、OAuth 流程、复杂组件。"],
        ["后端 handler 测试", "backend/internal/handler/**/*_test.go", "HTTP 参数、鉴权、错误返回、回调处理。"],
        ["后端 service 测试", "backend/internal/service/**/*_test.go", "账号调度、计费、OAuth、支付履约、Ops。"],
        ["repository/集成测试", "backend/internal/repository/**/*_test.go", "数据库查询、缓存、迁移、排序分页。"],
        ["Ent schema 测试", "backend/ent/schema/*_test.go", "字段、索引、关系、迁移兼容性。"],
    ], [3.2, 5.0, 7.0])
    add_para(doc, "建议补强方向：为“一个 API Key 为什么不可用”“一次请求为什么选中某账号”“一次支付为什么没有履约”这三类生产高频问题补充可观测性和自动化测试。")

    add_heading(doc, "附录 A：核心源码文件索引", 1)
    key_files = [
        "frontend/src/router/index.ts",
        "frontend/src/api/client.ts",
        "frontend/src/stores/auth.ts",
        "frontend/src/views/admin/AccountsView.vue",
        "frontend/src/views/admin/ops/OpsDashboard.vue",
        "frontend/src/views/user/KeysView.vue",
        "backend/cmd/server/main.go",
        "backend/internal/server/router.go",
        "backend/internal/server/routes/gateway.go",
        "backend/internal/server/routes/admin.go",
        "backend/internal/service/gateway_service.go",
        "backend/internal/service/openai_gateway_service.go",
        "backend/internal/service/account_service.go",
        "backend/internal/service/api_key_service.go",
        "backend/internal/service/payment_service.go",
        "backend/internal/repository/ent.go",
        "backend/ent/schema/account.go",
        "backend/ent/schema/api_key.go",
        "backend/ent/schema/usage_log.go",
    ]
    add_table(doc, ["文件", "用途"], [[f, "建议优先阅读的核心文件"] for f in key_files], [7.0, 8.0])

    add_heading(doc, "附录 B：术语表", 1)
    add_table(doc, ["术语", "解释"], [
        ["TS/TypeScript", "带类型的 JavaScript，用来减少前端字段和参数错误。"],
        ["Vue", "前端页面框架，用组件构建界面。"],
        ["Pinia", "Vue 的全局状态管理工具，类似前端的小型内存数据库。"],
        ["Axios", "前端发 HTTP 请求的库。"],
        ["Go", "后端编程语言，编译成一个可执行文件。"],
        ["Gin", "Go 的 Web 框架，负责 HTTP 路由和请求响应。"],
        ["Ent", "Go 的 ORM，用代码定义数据库表并生成查询代码。"],
        ["Middleware", "中间件，请求进入业务前的一层检查或增强。"],
        ["Handler", "HTTP 处理器，负责读请求和写响应。"],
        ["Service", "业务服务，承载核心规则。"],
        ["Repository", "数据访问层，负责数据库、Redis、外部服务。"],
        ["OAuth", "授权登录/授权账号的标准流程，例如 OpenAI/Codex 账号授权。"],
        ["Failover", "失败转移，一个账号失败时尝试换另一个可用账号。"],
        ["SSE/Stream", "流式响应，模型一边生成一边返回。"],
    ], [4.0, 11.0])

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = write_manual()
    print(path)
