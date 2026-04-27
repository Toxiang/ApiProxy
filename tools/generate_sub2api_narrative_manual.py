from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
IMG_DIR = OUT_DIR / "narrative_images"
DOCX_PATH = OUT_DIR / "Sub2API_专业版技术说明书_叙事版.docx"

ACCENT = RGBColor(32, 80, 129)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(89, 104, 124)
LIGHT_BLUE = "EAF3FF"
LIGHT_GREEN = "EAF8F0"
LIGHT_AMBER = "FFF6E5"
LIGHT_RED = "FFF0F0"
LIGHT_GRAY = "F6F8FB"


def find_font() -> str | None:
    for item in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]:
        if Path(item).exists():
            return item
    return None


FONT_PATH = find_font()


def image_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, width: int, font_obj) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            test = line + ch
            if draw.textbbox((0, 0), test, font=font_obj)[2] <= width:
                line = test
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
        if raw == "":
            lines.append("")
    return lines


def draw_wrapped(draw: ImageDraw.ImageDraw, xy, text: str, width: int, fill: str, font_obj, gap: int = 8) -> int:
    x, y = xy
    for line in wrap_text(draw, text, width, font_obj):
        draw.text((x, y), line, fill=fill, font=font_obj)
        y += getattr(font_obj, "size", 20) + gap
    return y


def rounded(draw: ImageDraw.ImageDraw, xy, fill: str, outline: str, radius: int = 24, width: int = 3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw: ImageDraw.ImageDraw, start, end, fill="#52627A", width=4):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    dx, dy = ex - sx, ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx > 0 else -1
        pts = [(ex, ey), (ex - 16 * sign, ey - 9), (ex - 16 * sign, ey + 9)]
    else:
        sign = 1 if dy > 0 else -1
        pts = [(ex, ey), (ex - 9, ey - 16 * sign), (ex + 9, ey - 16 * sign)]
    draw.polygon(pts, fill=fill)


def save_canvas(path: Path, title: str, subtitle: str, boxes: list[dict], links: list[tuple[int, int, str]]):
    img = Image.new("RGB", (1800, 1040), "#FBFCFE")
    d = ImageDraw.Draw(img)
    d.text((70, 42), title, fill="#172033", font=image_font(44))
    if subtitle:
        draw_wrapped(d, (72, 100), subtitle, 1500, "#64748B", image_font(24), 6)
    centers = []
    for box in boxes:
        x, y, w, h = box["rect"]
        centers.append((x + w // 2, y + h // 2))
    for a, b, label in links:
        sx, sy = centers[a]
        ex, ey = centers[b]
        arrow(d, (sx, sy), (ex, ey), fill="#7A879A")
        if label:
            mx, my = (sx + ex) // 2, (sy + ey) // 2
            d.rounded_rectangle((mx - 78, my - 22, mx + 78, my + 22), radius=14, fill="#FFFFFF", outline="#CBD5E1", width=2)
            d.text((mx - 62, my - 14), label, fill="#475569", font=image_font(18))
    for box in boxes:
        x, y, w, h = box["rect"]
        rounded(d, (x, y, x + w, y + h), box.get("fill", "#EEF4FF"), box.get("outline", "#4267D5"), 26, 3)
        d.text((x + 26, y + 22), box["title"], fill="#172033", font=image_font(27))
        draw_wrapped(d, (x + 26, y + 70), box["body"], w - 52, "#3E4A5F", image_font(22), 6)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_images() -> dict[str, Path]:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    images: dict[str, Path] = {}

    images["overview"] = IMG_DIR / "01_overview.png"
    save_canvas(
        images["overview"],
        "Sub2API 产品级架构图",
        "从用户入口到上游模型服务的完整分层。新版文档以这张图为主线，而不是把源码文件逐个罗列。",
        [
            {"rect": (70, 210, 280, 185), "title": "用户与客户端", "body": "浏览器后台\nCodex / Claude Code\nOpenAI / Gemini SDK", "fill": "#FFFFFF", "outline": "#94A3B8"},
            {"rect": (430, 145, 310, 210), "title": "Vue 前端", "body": "页面路由\nPinia 状态\nAxios API 封装\n支付与 OAuth 回调", "fill": "#EAF8F0", "outline": "#2F9E66"},
            {"rect": (430, 500, 330, 260), "title": "Go 后端网关", "body": "Gin 路由\n鉴权中间件\nService 业务层\n协议转换与计费", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (880, 175, 300, 190), "title": "PostgreSQL", "body": "用户、API Key\n账号、分组、订单\n用量与审计日志", "fill": "#FFF6E5", "outline": "#D99728"},
            {"rect": (880, 485, 300, 190), "title": "Redis", "body": "登录刷新 token\n限流计数\n调度缓存\n并发锁", "fill": "#FFF0F0", "outline": "#D44D4D"},
            {"rect": (1320, 240, 320, 260), "title": "上游 AI", "body": "Anthropic\nOpenAI / Codex\nGemini\nAntigravity", "fill": "#F1EDFF", "outline": "#7B61D1"},
            {"rect": (1320, 650, 320, 180), "title": "支付与外部服务", "body": "EasyPay / 支付宝 / 微信\nStripe\nS3 备份", "fill": "#EAF9FB", "outline": "#25A2B8"},
        ],
        [(0, 1, "访问"), (1, 2, "API"), (2, 3, "持久化"), (2, 4, "缓存"), (2, 5, "转发"), (2, 6, "支付")],
    )

    images["request"] = IMG_DIR / "02_request_flow.png"
    save_canvas(
        images["request"],
        "一次模型请求的真实处理链路",
        "这张图解释 /v1/responses、/v1/messages、/chat/completions 进入后端后怎样被鉴权、调度、转换、转发和计费。",
        [
            {"rect": (60, 230, 240, 190), "title": "1. 入站请求", "body": "用户携带系统 API Key\n请求模型接口", "fill": "#FFFFFF", "outline": "#94A3B8"},
            {"rect": (360, 230, 250, 190), "title": "2. 中间件", "body": "请求大小\nRequest ID\nAPI Key 鉴权\n分组检查", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (670, 230, 260, 190), "title": "3. 业务校验", "body": "余额/订阅\n用户并发\n分组模型限制", "fill": "#EAF8F0", "outline": "#2F9E66"},
            {"rect": (990, 230, 260, 190), "title": "4. 账号调度", "body": "选择可用账号\n检查限流\n粘性会话", "fill": "#FFF6E5", "outline": "#D99728"},
            {"rect": (1310, 230, 280, 190), "title": "5. 协议转换", "body": "Claude/OpenAI/Gemini\n请求体修正\n模型映射", "fill": "#F1EDFF", "outline": "#7B61D1"},
            {"rect": (690, 620, 300, 185), "title": "6. 上游响应", "body": "SSE 或 JSON 返回\n解析 usage\n写用量与运维日志", "fill": "#EAF9FB", "outline": "#25A2B8"},
        ],
        [(0, 1, ""), (1, 2, ""), (2, 3, ""), (3, 4, ""), (4, 5, ""), (5, 0, "返回")],
    )

    images["oauth"] = IMG_DIR / "03_openai_oauth_codex.png"
    save_canvas(
        images["oauth"],
        "OpenAI OAuth 代理 Codex 的完整过程",
        "重点理解：系统保存的是 OpenAI OAuth 账号凭证，用户调用时仍然只用本系统 API Key。",
        [
            {"rect": (65, 190, 290, 190), "title": "管理员添加账号", "body": "前端请求\n/api/v1/admin/openai/generate-auth-url", "fill": "#EAF8F0", "outline": "#2F9E66"},
            {"rect": (420, 190, 300, 190), "title": "生成授权参数", "body": "state\nPKCE verifier/challenge\nsession_id\n可选代理", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (785, 190, 300, 190), "title": "OpenAI 授权", "body": "auth.openai.com\nCodex CLI client_id\noffline_access", "fill": "#FFF6E5", "outline": "#D99728"},
            {"rect": (1150, 190, 315, 190), "title": "code 换 token", "body": "POST /oauth/token\naccess_token\nrefresh_token\nid_token", "fill": "#F1EDFF", "outline": "#7B61D1"},
            {"rect": (1150, 560, 315, 190), "title": "保存账号凭证", "body": "Account.credentials JSON\nchatgpt_account_id\nexpires_at", "fill": "#FFF0F0", "outline": "#D44D4D"},
            {"rect": (785, 560, 300, 190), "title": "网关取 token", "body": "缓存优先\n快过期则刷新\n分布式锁防并发刷新", "fill": "#EAF9FB", "outline": "#25A2B8"},
            {"rect": (420, 560, 300, 190), "title": "转换为 Codex 请求", "body": "store=false\nstream=true\n修正 model/tools/input", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (65, 560, 290, 190), "title": "请求 ChatGPT 内部 API", "body": "chatgpt.com/backend-api/codex/responses\nBearer access_token", "fill": "#FFFFFF", "outline": "#94A3B8"},
        ],
        [(0, 1, ""), (1, 2, ""), (2, 3, ""), (3, 4, ""), (4, 5, ""), (5, 6, ""), (6, 7, "")],
    )

    images["frontend"] = IMG_DIR / "04_frontend_flow.png"
    save_canvas(
        images["frontend"],
        "前端页面如何组织",
        "前端不是业务核心，但它决定管理员和用户如何安全地触发后端能力。",
        [
            {"rect": (90, 200, 320, 200), "title": "路由 router", "body": "公开页\n用户页\n管理员页\n支付页\nOAuth 回调页", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (520, 200, 320, 200), "title": "页面 views", "body": "Dashboard\nKeys / Usage\nAccounts / Groups\nOps / Payment", "fill": "#EAF8F0", "outline": "#2F9E66"},
            {"rect": (950, 200, 320, 200), "title": "组件 components", "body": "弹窗\n表格\n图表\n布局\n账号编辑器", "fill": "#FFF6E5", "outline": "#D99728"},
            {"rect": (520, 560, 320, 200), "title": "状态 stores", "body": "auth\napp\npayment\nsubscriptions\nadminSettings", "fill": "#F1EDFF", "outline": "#7B61D1"},
            {"rect": (950, 560, 320, 200), "title": "API 封装", "body": "api/client.ts\napi/auth.ts\napi/admin/*\n错误与 token 刷新", "fill": "#EAF9FB", "outline": "#25A2B8"},
        ],
        [(0, 1, "进入"), (1, 2, "复用"), (1, 3, "读写"), (3, 4, "请求"), (4, 1, "返回")],
    )

    images["backend"] = IMG_DIR / "05_backend_layers.png"
    save_canvas(
        images["backend"],
        "后端分层与职责边界",
        "读 Go 代码时不要从 ent 生成文件开始，而要顺着 HTTP 路由到 handler、service、repository 的链路看。",
        [
            {"rect": (120, 150, 330, 185), "title": "cmd/server", "body": "加载配置\n依赖注入\n启动 Gin\n优雅关闭", "fill": "#FFFFFF", "outline": "#94A3B8"},
            {"rect": (560, 150, 330, 185), "title": "server/routes", "body": "注册 URL\n挂载中间件\n区分 auth/user/admin/gateway", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (1000, 150, 330, 185), "title": "handler", "body": "读请求参数\n检查格式\n返回统一响应\n不承载复杂业务", "fill": "#EAF8F0", "outline": "#2F9E66"},
            {"rect": (1000, 520, 330, 185), "title": "service", "body": "核心业务规则\n账号调度\nOAuth 刷新\n计费和支付履约", "fill": "#FFF6E5", "outline": "#D99728"},
            {"rect": (560, 520, 330, 185), "title": "repository", "body": "数据库访问\nRedis 缓存\n外部 HTTP 客户端", "fill": "#F1EDFF", "outline": "#7B61D1"},
            {"rect": (120, 520, 330, 185), "title": "ent/schema", "body": "数据表定义\n迁移基础\n生成类型安全代码", "fill": "#EAF9FB", "outline": "#25A2B8"},
        ],
        [(0, 1, ""), (1, 2, ""), (2, 3, ""), (3, 4, ""), (4, 5, "")],
    )

    images["data"] = IMG_DIR / "06_data_model.png"
    save_canvas(
        images["data"],
        "核心数据关系",
        "平台能否稳定运行，取决于用户、API Key、分组、账号、用量、订单这些概念是否分清。",
        [
            {"rect": (90, 240, 260, 175), "title": "User", "body": "登录身份\n余额/角色\n并发额度", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (470, 240, 260, 175), "title": "APIKey", "body": "sk- 凭证\n额度/IP 限制\n绑定分组", "fill": "#EAF8F0", "outline": "#2F9E66"},
            {"rect": (850, 240, 260, 175), "title": "Group", "body": "平台类型\n模型策略\n倍率与限制", "fill": "#FFF6E5", "outline": "#D99728"},
            {"rect": (1230, 240, 280, 175), "title": "Account", "body": "OAuth/API Key\n代理/并发\n上游凭证", "fill": "#F1EDFF", "outline": "#7B61D1"},
            {"rect": (470, 610, 260, 175), "title": "PaymentOrder", "body": "支付金额\n渠道/状态\n履约结果", "fill": "#FFF0F0", "outline": "#D44D4D"},
            {"rect": (850, 610, 260, 175), "title": "Subscription", "body": "套餐权益\n有效期\n周期额度", "fill": "#F8F0E3", "outline": "#B98A4A"},
            {"rect": (1230, 610, 280, 175), "title": "UsageLog", "body": "模型请求\ntoken/费用\n延迟与错误", "fill": "#EAF9FB", "outline": "#25A2B8"},
        ],
        [(0, 1, "拥有"), (1, 2, "属于"), (2, 3, "调度"), (1, 6, "产生日志"), (0, 4, "购买"), (4, 5, "履约"), (5, 2, "授权")],
    )

    images["payment"] = IMG_DIR / "07_payment_flow.png"
    save_canvas(
        images["payment"],
        "支付、订阅与余额履约",
        "支付模块不只是收钱，还要保证回调验签、幂等履约、订单状态和用户权益一致。",
        [
            {"rect": (70, 250, 260, 180), "title": "选择套餐", "body": "用户在 /purchase\n读取可见套餐和渠道", "fill": "#EAF8F0", "outline": "#2F9E66"},
            {"rect": (420, 250, 260, 180), "title": "创建订单", "body": "生成订单号\n锁定金额和套餐\n等待支付", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (770, 250, 260, 180), "title": "第三方支付", "body": "EasyPay\n支付宝/微信\nStripe", "fill": "#FFF6E5", "outline": "#D99728"},
            {"rect": (1120, 250, 260, 180), "title": "回调验签", "body": "确认来源\n确认金额\n避免重复回调", "fill": "#FFF0F0", "outline": "#D44D4D"},
            {"rect": (770, 610, 260, 180), "title": "订单履约", "body": "开通订阅\n充值余额\n写审计日志", "fill": "#F1EDFF", "outline": "#7B61D1"},
            {"rect": (1120, 610, 260, 180), "title": "用户查询", "body": "订单列表\n支付结果页\n管理员对账", "fill": "#EAF9FB", "outline": "#25A2B8"},
        ],
        [(0, 1, ""), (1, 2, ""), (2, 3, ""), (3, 4, ""), (4, 5, "")],
    )

    images["ops"] = IMG_DIR / "08_ops_loop.png"
    save_canvas(
        images["ops"],
        "运维闭环：发现、定位、恢复",
        "网关类系统最重要的是可观测性。否则请求失败时只能猜是用户、账号、上游还是配置问题。",
        [
            {"rect": (100, 210, 300, 190), "title": "实时指标", "body": "QPS/TPS\n延迟\n并发槽位\n账号可用性", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (520, 210, 300, 190), "title": "错误采集", "body": "请求错误\n上游错误\n系统日志\n重试记录", "fill": "#FFF0F0", "outline": "#D44D4D"},
            {"rect": (940, 210, 300, 190), "title": "告警规则", "body": "阈值\n静默\n邮件通知\n事件状态", "fill": "#FFF6E5", "outline": "#D99728"},
            {"rect": (520, 580, 300, 190), "title": "人工处置", "body": "查看请求详情\n重试错误请求\n清理账号限流\n调整运行配置", "fill": "#EAF8F0", "outline": "#2F9E66"},
            {"rect": (940, 580, 300, 190), "title": "恢复验证", "body": "定时测试\n通道监控\n用量趋势\n系统日志健康", "fill": "#EAF9FB", "outline": "#25A2B8"},
        ],
        [(0, 1, ""), (1, 2, ""), (2, 3, ""), (3, 4, ""), (4, 0, "")],
    )

    images["deploy"] = IMG_DIR / "09_deployment_storage.png"
    save_canvas(
        images["deploy"],
        "部署与数据持久化",
        "回答你前面关心的问题：Docker 部署时数据要看 volume 绑定，PostgreSQL 数据目录不一定是容器里的 /var/lib/postgresql。",
        [
            {"rect": (100, 210, 330, 210), "title": "apiproxy 容器", "body": "运行 Go 服务\n对外 18080 -> 8080\n读取配置和静态前端", "fill": "#EAF3FF", "outline": "#4267D5"},
            {"rect": (540, 210, 330, 210), "title": "postgres 容器", "body": "真正数据通常在\n/var/lib/postgresql/data\n或 compose volume", "fill": "#FFF6E5", "outline": "#D99728"},
            {"rect": (980, 210, 330, 210), "title": "redis 容器", "body": "缓存、锁、限流\n是否持久化取决于 compose 配置", "fill": "#FFF0F0", "outline": "#D44D4D"},
            {"rect": (540, 580, 330, 210), "title": "宿主机目录", "body": "deploy/data\n或 docker volume\n才是重启后能保留的位置", "fill": "#EAF8F0", "outline": "#2F9E66"},
            {"rect": (980, 580, 330, 210), "title": "备份/恢复", "body": "后台备份\nS3 配置\n数据管理 agent", "fill": "#EAF9FB", "outline": "#25A2B8"},
        ],
        [(0, 1, "连接"), (0, 2, "连接"), (1, 3, "volume"), (2, 3, "volume"), (3, 4, "备份")],
    )

    return images


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_width(cell, cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(int(cm * 567)))
    width.set(qn("w:type"), "dxa")


def style_run(run, size=10.5, bold=False, color: RGBColor = TEXT, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def shade_paragraph(p, fill: str):
    p_pr = p._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_para(doc: Document, text: str = "", size=10.5, color: RGBColor = TEXT, bold=False, space_after=5):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.line_spacing = 1.28
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(0.0)
    r = p.add_run(text)
    style_run(r, size=size, color=color, bold=bold)
    return p


def add_heading(doc: Document, text: str, level=1):
    style_name = f"Heading {min(max(level, 1), 3)}"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    r = p.add_run(text)
    if level == 1:
        style_run(r, size=17, bold=True, color=ACCENT)
    elif level == 2:
        style_run(r, size=13.2, bold=True, color=RGBColor(45, 75, 115))
    else:
        style_run(r, size=11.5, bold=True, color=RGBColor(65, 85, 115))
    return p


def add_small_caps(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=8.8, bold=True, color=MUTED)
    return p


def add_code(doc: Document, lines: Iterable[str]):
    p = doc.add_paragraph()
    shade_paragraph(p, "F3F6FA")
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        style_run(r, size=9.2, color=RGBColor(43, 56, 72), font_name="Consolas")


def add_callout(doc: Document, title: str, body: str, fill=LIGHT_BLUE):
    p = doc.add_paragraph()
    shade_paragraph(p, fill)
    p.paragraph_format.left_indent = Cm(0.18)
    p.paragraph_format.right_indent = Cm(0.18)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    style_run(r, size=11.5, bold=True, color=ACCENT)
    p2 = doc.add_paragraph()
    shade_paragraph(p2, fill)
    p2.paragraph_format.left_indent = Cm(0.18)
    p2.paragraph_format.right_indent = Cm(0.18)
    p2.paragraph_format.line_spacing = 1.24
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(body)
    style_run(r2, size=10.2, color=TEXT)


def add_card_grid(doc: Document, cards: list[tuple[str, str, str]], cols=2):
    for title, body, fill in cards:
        p = doc.add_paragraph()
        shade_paragraph(p, fill)
        p.paragraph_format.left_indent = Cm(0.18)
        p.paragraph_format.right_indent = Cm(0.18)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title)
        style_run(r, size=10.8, bold=True, color=ACCENT)
        p2 = doc.add_paragraph()
        shade_paragraph(p2, fill)
        p2.paragraph_format.left_indent = Cm(0.18)
        p2.paragraph_format.right_indent = Cm(0.18)
        p2.paragraph_format.line_spacing = 1.22
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(body)
        style_run(r2, size=9.8, color=TEXT)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    header_line = " / ".join(headers)
    hp = doc.add_paragraph()
    shade_paragraph(hp, "DCEBFA")
    hp.paragraph_format.left_indent = Cm(0.18)
    hp.paragraph_format.space_before = Pt(5)
    hp.paragraph_format.space_after = Pt(3)
    hr = hp.add_run(header_line)
    style_run(hr, size=9.5, bold=True, color=ACCENT)
    for idx, row in enumerate(rows, start=1):
        fill = "FFFFFF" if idx % 2 else "F7FAFE"
        p = doc.add_paragraph()
        shade_paragraph(p, fill)
        p.paragraph_format.left_indent = Cm(0.18)
        p.paragraph_format.right_indent = Cm(0.18)
        p.paragraph_format.line_spacing = 1.22
        p.paragraph_format.space_after = Pt(5)
        for i, text in enumerate(row):
            if i:
                p.add_run("\n")
            label = headers[i] if i < len(headers) else f"字段{i + 1}"
            r_label = p.add_run(f"{label}：")
            style_run(r_label, size=9.2, bold=True, color=ACCENT)
            r_text = p.add_run(text)
            style_run(r_text, size=9.2, color=TEXT)


def add_picture(doc: Document, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.2))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    style_run(r, size=9, color=MUTED)
    c.paragraph_format.space_after = Pt(8)


def configure_doc(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.75)
    sec.bottom_margin = Cm(1.65)
    sec.left_margin = Cm(1.75)
    sec.right_margin = Cm(1.75)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [
        ("Heading 1", 17, ACCENT),
        ("Heading 2", 13.2, RGBColor(45, 75, 115)),
        ("Heading 3", 11.5, RGBColor(65, 85, 115)),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
    if "DocMeta" not in styles:
        meta = styles.add_style("DocMeta", WD_STYLE_TYPE.PARAGRAPH)
        meta.font.name = "Microsoft YaHei"
        meta._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        meta.font.size = Pt(9)
        meta.font.color.rgb = MUTED
        meta.paragraph_format.line_spacing = 1.18
        meta.paragraph_format.space_after = Pt(2)
    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = header.add_run("Sub2API Technical Manual")
        style_run(hr, size=8.5, color=MUTED)
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = footer.add_run("Sub2API 专业版技术说明书 | 叙事版")
        style_run(rr, size=8.5, color=MUTED)


def cover(doc: Document):
    add_para(doc, "", size=8, space_after=24)
    add_small_caps(doc, "TECHNICAL MANUAL / ARCHITECTURE AND OPERATIONS")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sub2API\n专业版技术说明书")
    style_run(r, size=30, bold=True, color=ACCENT)
    p.paragraph_format.space_after = Pt(14)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("面向非 TS / Go 背景读者的功能链路、代码解释、设计取舍与运维说明")
    style_run(r2, size=13, color=MUTED)
    add_para(doc, "", size=8, space_after=8)
    add_card_grid(doc, [
        ("文档定位", "架构说明 + 功能手册 + 代码导读 + 运维排障指南。", LIGHT_BLUE),
        ("适用读者", "项目维护者、部署人员、管理员、准备二次开发的工程师。", LIGHT_GREEN),
        ("阅读方式", "先看图和每章第一句话，再看流程，最后按代码索引进入源码。", LIGHT_AMBER),
    ])
    add_callout(
        doc,
        "这版文档的写法",
        "不再把文件列表当正文，而是像排查一个真实问题那样讲：先说明业务目标，再画流程，再解释关键代码，最后给出为什么这样设计、优缺点和可优化方向。",
        fill=LIGHT_BLUE,
    )
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n代码库：Z:\\Trans\\sub2api")
    style_run(r3, size=10, color=MUTED)
    doc.add_page_break()


def add_outline(doc: Document):
    add_heading(doc, "目录与阅读导航", 1)
    add_para(doc, "本文按产品级技术手册组织，正文不依赖文件清单展开。每章都遵循同一结构：先用一句话说明，再给流程图或卡片，接着解释关键代码和设计取舍。")
    add_simple_table(doc, ["章节", "阅读目标"], [
        ["1-2 系统和请求链路", "先建立整体地图，知道一次模型请求从哪里进、如何调度、怎么返回。"],
        ["3-4 前端和后端分层", "理解 Vue 前端和 Go 后端各自负责什么，以及读代码的正确入口。"],
        ["5-7 认证、账号池、OpenAI OAuth Codex", "理解用户登录、本系统 API Key、上游 OAuth 账号之间的区别。"],
        ["8-10 协议兼容、调度、支付计费", "理解平台为什么能兼容多客户端、多上游、多套餐。"],
        ["11-12 运维和部署", "理解怎么排查线上问题，以及 Docker 部署数据保存在哪里。"],
        ["13-15 选型、代码阅读、优化路线", "理解技术栈选择、后续维护方法和改进方向。"],
        ["附录", "只保留关键代码索引和术语解释，避免正文退化成文件列表。"],
    ], [4.2, 10.5])
    add_callout(doc, "排版原则", "本版参考成熟技术文档常见做法：短段落、描述性标题、明显层级、稳定字体、图文配合、代码路径单独成块、文件索引放附录。", fill=LIGHT_GRAY)
    doc.add_page_break()


def section_intro(doc: Document, num: str, title: str, summary: str, code_refs: list[str] | None = None):
    add_heading(doc, f"{num}. {title}", 1)
    add_callout(doc, "先用一句话讲明白", summary, fill=LIGHT_BLUE)
    if code_refs:
        add_code(doc, ["本章重点代码：", *[f"- {ref}" for ref in code_refs]])


def write_manual() -> Path:
    images = make_images()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    cover(doc)
    add_outline(doc)

    add_heading(doc, "阅读路径", 1)
    add_para(doc, "这份文档按“先理解系统，再理解代码，再理解运维”的顺序组织。你不需要先懂 TypeScript 或 Go，可以先顺着流程图理解每个功能为什么存在，再按代码路径回到源码。")
    add_card_grid(doc, [
        ("第一遍：看业务", "读第 1 到第 4 章，知道系统给谁用、有哪些入口、一次请求如何流转。", LIGHT_GREEN),
        ("第二遍：看核心", "读第 5 到第 10 章，理解登录、API Key、账号池、OAuth、调度、计费这些核心能力。", LIGHT_BLUE),
        ("第三遍：看生产", "读第 11 到第 14 章，理解日志、运维、部署、数据保存和排障方法。", LIGHT_AMBER),
        ("最后：看代码", "读附录和每章的代码索引，再进入具体文件。文件列表只作为索引，不再喧宾夺主。", LIGHT_GRAY),
    ])
    add_picture(doc, images["overview"], "图 1：Sub2API 的产品级架构。")

    section_intro(
        doc,
        "1",
        "系统到底是什么",
        "Sub2API 是一个多用户、多账号、多上游模型的 API 网关和管理平台：前端负责让用户和管理员配置资源，后端负责鉴权、调度、协议转换、计费、支付和运维监控。",
        ["frontend/src/router/index.ts", "backend/internal/server/router.go", "backend/internal/server/routes/gateway.go"],
    )
    add_para(doc, "从使用者角度看，它像一个统一的模型服务入口。用户拿到本系统的 API Key 后，可以用类似 OpenAI、Claude 或 Gemini 的接口发请求。系统内部会根据 API Key 所属分组选择合适的上游账号，并在必要时把请求转换成上游能理解的格式。")
    add_para(doc, "从管理员角度看，它是一个账号池和商业化管理后台。管理员可以接入 OpenAI OAuth、Claude、Gemini、Antigravity 等账号，配置分组、模型价格、用户额度、支付套餐、代理、监控和告警。")
    add_card_grid(doc, [
        ("前端的职责", "提供登录、用户中心、API Key 管理、支付、管理员后台、账号管理、运维监控等界面。它不直接调用上游 AI，而是调用后端接口。", LIGHT_GREEN),
        ("后端的职责", "后端是核心：验证身份、判断权限、选择上游账号、转换请求格式、转发请求、记录用量、处理支付、维护 OAuth token。", LIGHT_BLUE),
        ("数据库的职责", "保存长期业务数据，例如用户、API Key、账号凭证、分组、用量日志、订单、订阅、系统配置。", LIGHT_AMBER),
        ("Redis 的职责", "保存短期高频状态，例如限流计数、并发锁、token 缓存、调度缓存、登录刷新 token 状态。", LIGHT_RED),
    ])
    add_callout(doc, "为什么不要从文件列表开始读", "这个项目文件很多，如果一开始就看目录，很容易迷路。正确方式是先抓住 5 条主线：用户认证、API Key、账号池、网关请求、计费支付。知道这些概念后，再看文件才有方向。", fill=LIGHT_AMBER)

    section_intro(
        doc,
        "2",
        "一次模型请求如何流转",
        "一次模型请求进来后，后端会先做本系统鉴权，再做余额和并发检查，然后调度账号，转换请求，发给上游，最后把响应和用量记录下来。",
        ["backend/internal/server/routes/gateway.go", "backend/internal/handler/openai_gateway_handler.go", "backend/internal/service/openai_gateway_service.go", "backend/internal/service/gateway_service.go"],
    )
    add_picture(doc, images["request"], "图 2：模型请求从入口到上游再返回的处理链路。")
    add_para(doc, "用户请求不会直接拿 OpenAI、Claude 或 Gemini 的真实密钥访问上游。用户只带本系统的 API Key，例如 `Authorization: Bearer sk-...`。这把 Key 首先在中间件里被识别，并且会绑定到用户、分组、订阅、限流和计费规则。")
    add_para(doc, "后端的路由设计让多个协议入口可以共用一套底层调度能力。例如 `/v1/messages` 可以兼容 Claude Messages，也可以在分组平台是 OpenAI 时转给 OpenAI 网关；`/v1/responses` 和 `/responses` 都能进入 Responses 风格处理；`/backend-api/codex/responses` 是面向 Codex 兼容场景的直达别名。")
    add_simple_table(doc, ["阶段", "做什么", "失败时会怎样"], [
        ["入站鉴权", "检查 API Key、用户状态、分组关系和请求体大小。", "返回 401、403 或请求体过大错误。"],
        ["业务校验", "检查余额、订阅权益、用户并发、分组模型限制。", "返回余额不足、订阅不可用或并发限制。"],
        ["账号调度", "从账号池里选择支持当前模型且状态可用的账号。", "没有账号时返回服务暂不可用，或进入 failover。"],
        ["协议转换", "把客户端协议转换成目标上游协议，修正模型名、tools、stream、参数。", "参数不兼容时返回 400 或上游错误。"],
        ["上游转发", "按账号类型选择 OAuth/API Key，带代理和必要请求头访问上游。", "429、5xx、超时等可能触发换账号。"],
        ["用量记录", "解析 token、模型、成本、请求延迟和错误信息。", "记录失败不应影响主请求返回，但会写日志。"],
    ], [3.0, 6.0, 6.0])

    section_intro(
        doc,
        "3",
        "前端如何工作",
        "前端是一个 Vue 3 管理后台：router 决定用户能访问哪些页面，stores 保存登录和全局状态，api 目录封装所有后端请求，views 和 components 组成实际页面。",
        ["frontend/src/router/index.ts", "frontend/src/api/client.ts", "frontend/src/stores/auth.ts", "frontend/src/views/admin/AccountsView.vue", "frontend/src/views/admin/ops/OpsDashboard.vue"],
    )
    add_picture(doc, images["frontend"], "图 3：前端页面、组件、状态、API 封装之间的关系。")
    add_para(doc, "如果你不熟悉 TypeScript，可以把前端理解成 4 层：路由层负责“去哪个页面”，页面层负责“显示和操作”，状态层负责“记住当前登录、设置、支付状态”，API 层负责“和后端说话”。")
    add_card_grid(doc, [
        ("router/index.ts", "定义 `/login`、`/dashboard`、`/admin/accounts`、`/purchase` 等页面，并用 navigation guard 判断是否登录、是否管理员、支付功能是否开启。", LIGHT_BLUE),
        ("api/client.ts", "统一 Axios 客户端。它负责给请求带 token、处理错误、刷新登录态，是前端和后端之间的主通道。", LIGHT_GREEN),
        ("stores/auth.ts", "保存当前用户、access token、角色、简单模式、待绑定 OAuth 状态等。页面刷新时会恢复登录状态。", LIGHT_AMBER),
        ("views/admin/*", "管理员功能页面，包括账号池、用户、分组、价格、代理、设置、运维监控、支付订单等。", LIGHT_GRAY),
        ("views/user/*", "普通用户页面，包括仪表盘、API Key、用量、购买套餐、订单、个人资料、通道状态等。", LIGHT_GRAY),
        ("components/*", "复用组件，例如表格、弹窗、图表、账号编辑器、支付组件、布局组件，避免每个页面重复造轮子。", LIGHT_GREEN),
    ])
    add_callout(doc, "前端设计取舍", "优点是页面入口清晰、功能覆盖完整、权限由路由统一控制。缺点是后台页面较多，部分复杂页面容易变成大文件。后续优化可以把账号编辑、OAuth 绑定、支付配置、Ops 图表拆成更小的业务组件。", fill=LIGHT_AMBER)

    section_intro(
        doc,
        "4",
        "后端如何启动和分层",
        "Go 后端从 cmd/server 启动，经依赖注入组装 handler、service、repository，再通过 routes 挂载 HTTP 接口。",
        ["backend/cmd/server/main.go", "backend/cmd/server/wire_gen.go", "backend/internal/server/router.go", "backend/internal/handler", "backend/internal/service", "backend/internal/repository"],
    )
    add_picture(doc, images["backend"], "图 4：后端从路由到业务再到数据访问的分层。")
    add_para(doc, "读后端代码时，不建议从 `ent` 生成目录开始。更好的路线是：先看 `server/router.go` 知道有哪些大类路由，再看 `server/routes/*.go` 找到具体 URL，接着看对应 handler，最后进入 service 读业务规则。repository 和 ent 只在需要确认数据保存时再深入。")
    add_simple_table(doc, ["层级", "它回答的问题", "新手怎么读"], [
        ["cmd/server", "程序如何启动、加载配置、连接数据库、启动后台任务。", "先看 main.go，再看 wire_gen.go 的依赖组装。"],
        ["routes", "一个 URL 会进入哪个 handler，会经过哪些中间件。", "查接口时先来这里，不要全局乱搜。"],
        ["handler", "HTTP 参数怎么读，错误怎么返回。", "handler 应该薄，复杂逻辑通常会转给 service。"],
        ["service", "真正业务规则在哪里。", "这是理解系统最重要的目录。"],
        ["repository", "数据和外部服务怎么访问。", "只在追踪保存、查询、Redis、HTTP 客户端时看。"],
        ["ent/schema", "数据库表结构怎么定义。", "只改 schema，不要手动改生成文件。"],
    ], [3.0, 6.0, 6.0])

    section_intro(
        doc,
        "5",
        "登录、用户与会话",
        "用户认证分两类：普通用户登录本系统，和管理员接入第三方 OAuth 账号。前者决定谁能使用平台，后者决定平台能代理哪些上游账号。",
        ["backend/internal/server/routes/auth.go", "backend/internal/handler/auth_handler.go", "backend/internal/service/auth_service.go", "frontend/src/api/auth.ts", "frontend/src/stores/auth.ts"],
    )
    add_para(doc, "普通登录使用本系统自己的 access token 和 refresh token。前端登录后把 refresh token 保存在 localStorage，后续请求通过 `api/client.ts` 携带 access token。access token 过期后，前端调用 `/api/v1/auth/refresh` 换新 token。")
    add_para(doc, "登录相关接口设置了限流，例如注册、登录、验证码、刷新 token、找回密码。这是为了防止暴力破解、验证码轰炸和 refresh 接口被刷。代码里 Redis 限流失败时选择 fail-close，也就是 Redis 出问题时高风险入口会倾向于拒绝，这是安全优先的设计。")
    add_card_grid(doc, [
        ("普通邮箱登录", "注册、登录、验证码、找回密码、重置密码、当前用户信息。适合最基础的账号体系。", LIGHT_BLUE),
        ("第三方登录", "LinuxDo、OIDC、微信 OAuth，可以登录、绑定当前用户、待创建账号、待绑定登录。", LIGHT_GREEN),
        ("2FA/TOTP", "支持二次验证，用户可以启用、禁用、获取验证方式。", LIGHT_AMBER),
        ("会话撤销", "支持 logout 和 revoke-all-sessions，便于用户主动退出和管理员风控。", LIGHT_RED),
    ])
    add_callout(doc, "设计原因", "普通登录和上游 OAuth 必须分开讲。普通登录解决“谁在使用 Sub2API”，OpenAI/Gemini/Antigravity OAuth 解决“Sub2API 用哪个上游账号去请求模型”。两者都是 OAuth/token，但业务含义完全不同。", fill=LIGHT_AMBER)

    section_intro(
        doc,
        "6",
        "API Key、分组和账号池",
        "API Key 是用户请求入口，分组是策略和价格边界，账号池是真正承载上游请求的资源池。",
        ["backend/internal/server/middleware/api_key_auth.go", "backend/internal/service/api_key_service.go", "backend/internal/service/account_service.go", "backend/internal/service/openai_account_scheduler.go", "backend/ent/schema/group.go", "backend/ent/schema/account.go"],
    )
    add_picture(doc, images["data"], "图 5：用户、API Key、分组、账号、订阅、用量之间的关系。")
    add_para(doc, "用户创建 API Key 后，实际请求模型时只需要使用这把 Key。Key 绑定到某个分组，分组决定平台类型、可用模型、计费倍率、是否要求 OAuth 账号、是否支持 Claude Code 或 OpenAI Messages 等能力。")
    add_para(doc, "账号池是管理员配置的上游账号集合。一个账号可以是 OAuth 类型，也可以是 API Key 类型。账号里保存上游凭证、代理、并发数、优先级、模型映射、透传开关、限流状态和错误状态。")
    add_simple_table(doc, ["概念", "一句话解释", "常见问题"], [
        ["User", "平台里的用户，拥有余额、角色、并发额度和订阅。", "用户没余额或订阅不可用时，请求会被拒绝。"],
        ["API Key", "用户给程序调用接口用的凭证。", "Key 状态、分组、IP 限制都会影响请求。"],
        ["Group", "把模型、价格、平台和策略组织在一起。", "模型不可用往往是分组策略问题。"],
        ["Account", "真正访问上游 AI 的账号或密钥。", "账号过期、限流、代理失败会导致上游错误。"],
        ["Channel", "模型定价和通道配置的抽象。", "计费异常通常要查通道价格和倍率。"],
    ], [2.6, 6.2, 6.2])
    add_callout(doc, "为什么设计得这么绕", "如果系统只给一个人用，可以直接保存一个 OpenAI Key。但这个项目是多用户、多套餐、多模型、多上游账号的平台，所以必须把“用户身份”“调用凭证”“业务套餐”“上游账号”拆开。这样才能做到隔离、限额、计费和 failover。", fill=LIGHT_BLUE)

    section_intro(
        doc,
        "7",
        "OpenAI OAuth 代理 Codex",
        "系统先把 ChatGPT/OpenAI 账号通过 OAuth 授权进来，保存 access_token 和 refresh_token；用户请求时，网关用这个 OAuth access_token 去请求 chatgpt.com/backend-api/codex/responses。",
        ["backend/internal/handler/admin/openai_oauth_handler.go", "backend/internal/service/openai_oauth_service.go", "backend/internal/pkg/openai/oauth.go", "backend/internal/service/openai_token_provider.go", "backend/internal/service/openai_codex_transform.go", "backend/internal/service/openai_gateway_service.go"],
    )
    add_picture(doc, images["oauth"], "图 6：OpenAI OAuth 账号接入和 Codex 代理请求。")
    add_para(doc, "这里最容易混淆，所以单独讲细一点。OpenAI OAuth 代理 Codex 不是让用户直接拿 OpenAI 的 token。用户仍然使用本系统 API Key。区别在于后端调度到 OpenAI OAuth 账号时，会从数据库读取这个账号保存的 access_token，然后代替用户去请求 ChatGPT/Codex 的内部接口。")
    add_heading(doc, "7.1 授权 URL 怎么生成", 2)
    add_para(doc, "管理员在账号管理页点击 OpenAI OAuth 添加账号，前端请求 `/api/v1/admin/openai/generate-auth-url`。后端生成 `state`、`code_verifier`、`code_challenge` 和 `session_id`。`state` 防止 CSRF，PKCE 防止授权码被截获后直接换 token，session_id 用来在内存里暂存本次授权流程。")
    add_code(doc, [
        "授权 URL 里包含的关键字段：",
        "response_type=code",
        "client_id=app_EMoamEEZ73f0CkXaXp7hrann",
        "scope=openid profile email offline_access",
        "code_challenge_method=S256",
        "id_token_add_organizations=true",
        "codex_cli_simplified_flow=true",
    ])
    add_heading(doc, "7.2 code 怎么换 token", 2)
    add_para(doc, "OpenAI 授权完成后会返回 code 和 state。后端先根据 session_id 找到之前的 OAuthSession，再用常量时间比较校验 state。校验通过后，后端用 code、code_verifier、redirect_uri 和 client_id 请求 `https://auth.openai.com/oauth/token`。")
    add_para(doc, "换回来的 tokenInfo 里包括 `access_token`、`refresh_token`、`id_token`、`expires_at`、`client_id`、email、ChatGPT 账号 ID、用户 ID、组织 ID、套餐类型等。后端会尝试解析 id_token，并用 ChatGPT backend-api 补全 plan_type、subscription_expires_at 等信息。")
    add_heading(doc, "7.3 token 存在哪里", 2)
    add_para(doc, "长期保存位置是数据库里的 OpenAI OAuth Account 的 `credentials` JSON。临时 OAuth session 只存在内存 SessionStore，默认 30 分钟过期。也就是说，Docker 部署时真正要持久化的是 PostgreSQL 数据，而不是某个容器内的 `/var/lib/postgresql` 空目录。")
    add_code(doc, [
        "Account.credentials 里常见字段：",
        "access_token, refresh_token, id_token, expires_at, client_id",
        "email, chatgpt_account_id, chatgpt_user_id, organization_id, plan_type",
        "subscription_expires_at, privacy_mode",
    ])
    add_heading(doc, "7.4 用户请求时怎么用这个 token", 2)
    add_para(doc, "当请求进入 OpenAI 网关，系统先调度一个 OpenAI OAuth 账号，然后通过 `OpenAITokenProvider.GetAccessToken()` 拿 access_token。它会先查 Redis/token cache。如果缓存没有，或者数据库里的 expires_at 快到期，就通过统一的 OAuthRefreshAPI 刷新。")
    add_para(doc, "刷新时用了本地锁、Redis 分布式锁、DB 重读和 invalid_grant 竞态恢复。这样做是为了防止多个请求同时使用同一个 refresh_token。OAuth refresh_token 有时会轮换，如果并发刷新不加锁，很容易一个请求成功、另一个请求把账号标成失效。")
    add_heading(doc, "7.5 请求体为什么要转换", 2)
    add_para(doc, "OAuth 账号走的是 `https://chatgpt.com/backend-api/codex/responses`，不是公开的 `https://api.openai.com/v1/responses`。ChatGPT/Codex internal API 对参数要求不同，所以系统会在 `openai_codex_transform.go` 里修正请求。")
    add_simple_table(doc, ["转换点", "为什么要做"], [
        ["模型名规范化", "把 gpt-5.3、codex 等别名映射到上游能识别的 Codex/GPT 模型。"],
        ["store=false", "OAuth Codex internal API 要求不存储，否则可能返回 Store must be set to false。"],
        ["stream=true", "Codex internal endpoint 主要按 SSE 流式返回。"],
        ["删除不支持参数", "temperature、top_p、presence_penalty、prompt_cache_retention 等可能被上游拒绝。"],
        ["functions 转 tools", "兼容旧 ChatCompletions 风格工具调用。"],
        ["system role 转 instructions", "internal API 对 input 里的 system role 支持不同。"],
        ["input 字符串转消息数组", "Codex endpoint 要求 input 是列表结构。"],
        ["session_id 隔离", "避免不同 API Key 或不同用户的 prompt_cache_key 碰撞。"],
    ], [4.0, 11.0])
    add_callout(doc, "这一章最重要的结论", "OpenAI OAuth 代理 Codex 的本质是：Sub2API 保存一组 ChatGPT/Codex 可用身份，用户请求进来后，系统把用户请求翻译成 Codex internal API 请求，再用账号池里的 OAuth access_token 转发。用户看见的是统一 API，后端承担 token 刷新、参数兼容、账号调度和用量记录。", fill=LIGHT_AMBER)

    section_intro(
        doc,
        "8",
        "多平台协议兼容",
        "系统同时兼容 Claude、OpenAI Responses、OpenAI Chat Completions、Gemini v1beta、Antigravity 等入口，核心是把不同客户端协议统一转到账号池和上游服务。",
        ["backend/internal/pkg/apicompat", "backend/internal/service/gateway_forward_as_chat_completions.go", "backend/internal/service/gateway_forward_as_responses.go", "backend/internal/service/openai_messages_dispatch.go", "backend/internal/handler/gateway_handler.go"],
    )
    add_para(doc, "协议兼容层让用户可以用熟悉的客户端接入。例如 Claude Code 可以走 `/v1/messages`，OpenAI SDK 可以走 `/v1/responses` 或 `/chat/completions`，Gemini SDK 可以走 `/v1beta/models/...`。后端不会把这些入口完全写成独立系统，而是尽量复用鉴权、调度、计费、日志这几套核心能力。")
    add_card_grid(doc, [
        ("Claude Messages", "用于 Claude Code 或 Anthropic 风格客户端。OpenAI 分组也可以接收 `/v1/messages` 并转成 OpenAI 上游请求。", LIGHT_BLUE),
        ("OpenAI Responses", "现代 OpenAI 风格请求入口，也是 Codex OAuth 代理的主要请求体形态。", LIGHT_GREEN),
        ("Chat Completions", "兼容旧 OpenAI SDK 或一些第三方工具，必要时转成 Responses。", LIGHT_AMBER),
        ("Gemini v1beta", "兼容 Gemini SDK/CLI 的模型列表和模型调用路径。", LIGHT_GRAY),
        ("Antigravity", "有专门路由，可强制平台为 antigravity，避免混合调度。", LIGHT_GRAY),
        ("图片生成", "OpenAI 平台下支持 images generations/edits，并在 Responses image_generation 场景做归一化。", LIGHT_GREEN),
    ])
    add_callout(doc, "设计代价", "兼容越多协议，转换边界越复杂。后续维护时要特别关注模型参数、工具调用、流式响应、错误格式和 usage 统计是否仍然一致。", fill=LIGHT_AMBER)

    section_intro(
        doc,
        "9",
        "账号调度、限流与失败切换",
        "账号调度的目标是把请求分配给最合适的上游账号，并在账号限流、过载、超时或 5xx 时尽量切换到其他可用账号。",
        ["backend/internal/service/openai_account_scheduler.go", "backend/internal/service/gateway_service.go", "backend/internal/service/rate_limit_service.go", "backend/internal/handler/failover_loop.go"],
    )
    add_para(doc, "调度不是简单轮询。系统会考虑账号状态、分组绑定、模型支持、并发槽位、临时不可调度状态、限流恢复时间、优先级、粘性会话、池模式重试等因素。")
    add_simple_table(doc, ["因素", "作用"], [
        ["状态 status", "错误、禁用、过期账号不应被调度。"],
        ["分组绑定", "API Key 所在分组只能使用允许的账号池。"],
        ["模型支持", "账号或通道必须支持用户请求的模型。"],
        ["并发槽位", "避免同一账号被太多流式请求占满。"],
        ["临时不可调度", "token 刷新失败、429、529、过载等会让账号暂时冷却。"],
        ["粘性会话", "带 previous_response_id 或 prompt_cache_key 的请求尽量落到同一账号。"],
        ["failover", "可重试错误发生时，排除失败账号，尝试下一个账号。"],
    ], [4.0, 11.0])
    add_callout(doc, "为什么这块复杂", "模型网关的稳定性不是靠一次请求成功，而是靠大量请求下仍然能把坏账号、限流账号、网络差账号隔离出去。复杂调度的收益是稳定，代价是排查时必须有足够日志解释“为什么选这个账号”。", fill=LIGHT_BLUE)

    section_intro(
        doc,
        "10",
        "支付、订阅、余额与用量",
        "支付模块负责收款，订阅模块负责权益，用量模块负责扣费和审计。三者合在一起才形成商业化闭环。",
        ["backend/internal/payment", "backend/internal/handler/payment_handler.go", "backend/internal/handler/payment_webhook_handler.go", "backend/internal/service/payment_service.go", "backend/internal/service/usage_billing.go", "frontend/src/views/user/PaymentView.vue", "frontend/src/views/admin/orders"],
    )
    add_picture(doc, images["payment"], "图 7：支付、回调、履约和订单查询流程。")
    add_para(doc, "用户购买套餐时，前端先展示可用支付方式和套餐。后端创建订单，第三方支付完成后回调本系统。系统必须验签、校验金额、保证幂等，然后才会给用户开通订阅或充值余额。")
    add_para(doc, "用量记录则发生在每次模型请求完成后。系统会根据模型、token、倍率、订阅权益、缓存命中、图片费用等计算成本，并写入 UsageLog。管理员和用户看到的用量、趋势、余额变化都来自这些记录或聚合结果。")
    add_callout(doc, "支付模块的安全重点", "支付回调可能重复、延迟甚至被伪造，所以不能只看前端支付结果页。必须以后端回调验签和订单状态机为准。订单履约要幂等，失败后要能重试或人工修复。", fill=LIGHT_RED)

    section_intro(
        doc,
        "11",
        "运维监控与问题排查",
        "Ops 模块把系统运行中的请求、错误、延迟、并发、账号状态、OpenAI token 统计和系统日志集中展示，帮助定位生产问题。",
        ["frontend/src/views/admin/ops/OpsDashboard.vue", "backend/internal/handler/admin/ops_dashboard_handler.go", "backend/internal/service/ops_service.go", "backend/internal/service/ops_openai_token_stats.go"],
    )
    add_picture(doc, images["ops"], "图 8：运维监控从发现到恢复的闭环。")
    add_para(doc, "网关问题一般分为四类：用户侧问题、配置侧问题、账号侧问题、上游侧问题。Ops 的价值就是把这四类问题拆开，避免所有失败都被笼统地看成“接口不通”。")
    add_simple_table(doc, ["现象", "优先检查"], [
        ["用户说 API 不通", "API Key 状态、分组、余额、订阅、IP 限制、请求体格式。"],
        ["只有某个模型不通", "分组模型限制、通道定价、账号可用模型、模型映射。"],
        ["OpenAI OAuth 请求失败", "access_token 是否过期、refresh 是否失败、代理是否可达、账号是否 codex_cli_only。"],
        ["大量 429/529", "账号限流快照、上游过载、并发配置、failover 是否生效。"],
        ["支付成功但没到账", "订单状态、回调日志、验签、履约记录、是否重复回调。"],
        ["Docker 访问不通", "端口映射、容器监听地址、反向代理、健康检查与实际 HTTP 协议。"],
    ], [5.0, 10.0])

    section_intro(
        doc,
        "12",
        "部署与数据持久化",
        "Docker 部署时，数据是否保存在本地取决于 compose 里的 volume 映射；数据库数据通常不是看容器里 `/var/lib/postgresql` 是否有文件，而要看 PostgreSQL 实际 PGDATA 和宿主机 volume。",
        ["deploy/docker-compose*.yml", "deploy/*.sh", "deploy/*.service", "Dockerfile", "backend/internal/setup"],
    )
    add_picture(doc, images["deploy"], "图 9：Docker 部署时容器、宿主机目录和数据保存关系。")
    add_para(doc, "你之前看到容器里的 `/var/lib/postgresql` 是空的，这不一定代表没有数据。很多 PostgreSQL 镜像实际数据目录是 `/var/lib/postgresql/data`，并且 Docker volume 可能被挂载到这个子目录。判断数据是否持久化，应当看 compose 中 postgres 服务的 volumes 配置，以及 `docker inspect` 里实际 Mounts。")
    add_para(doc, "OpenAI OAuth token、用户、API Key、账号配置、订单和用量都应保存在 PostgreSQL。Redis 保存的是缓存、锁、限流和部分会话状态，通常可以重建，但生产环境仍建议持久化或至少监控。")
    add_callout(doc, "部署建议", "生产环境优先固定一套部署方式，并写清楚：端口映射、数据目录、备份目录、.env 位置、数据库 volume、Redis volume、反向代理配置和恢复步骤。否则迁移服务器或排查故障时很容易误删或找不到数据。", fill=LIGHT_AMBER)

    section_intro(
        doc,
        "13",
        "技术选型为什么这样定",
        "Vue 3、TypeScript、Go、Gin、Ent、PostgreSQL、Redis、Docker Compose 不是堆技术名词，而是分别服务于后台交互、类型安全、网关性能、数据一致性和部署便利。",
    )
    add_simple_table(doc, ["技术", "为什么选", "代价和优化"], [
        ["Vue 3", "适合管理后台，组件化清晰，生态成熟。", "复杂页面要持续拆组件，避免大文件。"],
        ["TypeScript", "前端字段和接口更安全，减少拼错字段。", "需要维护类型，建议未来从 OpenAPI 生成类型。"],
        ["Pinia", "轻量管理登录、设置、支付、订阅等全局状态。", "store 过多时要划清职责边界。"],
        ["Vite", "开发启动快，构建现代前端体验好。", "低内存服务器构建可能 OOM，生产建议 CI 构建或增加 swap。"],
        ["Go", "单二进制、并发强、部署简单，适合 API 网关。", "错误处理显式，新手读代码要从调用链入手。"],
        ["Gin", "轻量高性能，路由和中间件直观。", "团队需要自律保持分层。"],
        ["Ent", "schema 定义清晰，生成类型安全查询代码。", "生成文件多，新手不要直接改生成代码。"],
        ["PostgreSQL", "强一致保存业务数据，支持 JSONB 和复杂查询。", "需要备份、迁移和慢查询治理。"],
        ["Redis", "适合缓存、锁、限流、调度状态。", "不是主数据源，故障要有降级策略。"],
        ["Docker Compose", "把 app、db、redis 编排在一起，部署门槛低。", "volume、端口、env 必须写清楚。"],
    ], [3.0, 6.0, 6.0])

    section_intro(
        doc,
        "14",
        "怎么继续读代码",
        "读代码不要全局搜索到哪里算哪里，要按问题类型走固定路线。",
    )
    add_simple_table(doc, ["你想弄懂什么", "推荐阅读顺序"], [
        ["登录为什么失败", "frontend/src/api/auth.ts -> backend/internal/server/routes/auth.go -> auth_handler.go -> auth_service.go"],
        ["API Key 为什么不能请求模型", "routes/gateway.go -> api_key_auth.go -> api_key_service.go -> group/account service -> usage/ops 日志"],
        ["OpenAI OAuth 怎么刷新", "openai_oauth_handler.go -> openai_oauth_service.go -> openai_token_provider.go -> oauth_refresh_api.go"],
        ["Codex 请求为什么被改写", "openai_gateway_handler.go -> openai_gateway_service.go -> openai_codex_transform.go"],
        ["账号为什么没被选中", "openai_account_scheduler.go -> gateway_service.go -> rate_limit_service.go -> account status/extra"],
        ["支付为什么没到账", "payment_handler.go -> payment_webhook_handler.go -> payment_service.go -> payment provider -> order/subscription service"],
        ["前端某个页面数据从哪来", "router/index.ts -> 对应 views/*.vue -> api/*.ts -> 后端 routes/handler/service"],
        ["Docker 数据在哪里", "deploy/docker-compose*.yml -> docker inspect Mounts -> PostgreSQL PGDATA -> backend config"],
    ], [5.0, 10.0])

    section_intro(
        doc,
        "15",
        "优缺点与后续优化路线",
        "当前系统能力很完整，但复杂度也已经不低。后续最值得做的是增强可解释性、契约文档、部署标准化和关键流程自动化测试。",
    )
    add_card_grid(doc, [
        ("优势 1：平台闭环完整", "用户、API Key、账号池、模型网关、计费、支付、订阅、运维监控都已经形成闭环。", LIGHT_GREEN),
        ("优势 2：上游兼容广", "支持多种客户端协议和多个上游平台，迁移成本低。", LIGHT_GREEN),
        ("优势 3：生产意识强", "有 token 自动刷新、failover、Ops、限流、并发控制和错误记录。", LIGHT_GREEN),
        ("短板 1：文档与配置解释不足", "很多能力已经有代码，但缺少“为什么这样配置”的用户级说明。", LIGHT_AMBER),
        ("短板 2：协议转换复杂", "OpenAI/Codex/Claude/Gemini 参数差异很大，需要持续测试防回归。", LIGHT_AMBER),
        ("短板 3：部署路径多", "Docker、裸机、agent、备份、S3 都支持，但如果文档不统一，用户容易部署混乱。", LIGHT_RED),
    ])
    add_simple_table(doc, ["优化方向", "具体建议"], [
        ["接口契约", "补充 OpenAPI 文档或从后端生成前端类型，减少 TS/Go 字段漂移。"],
        ["调度解释", "后台增加“为什么选择或不选择某账号”的诊断面板。"],
        ["OAuth 可视化", "展示 token 到期时间、最近刷新结果、代理连通性、是否需要重新授权。"],
        ["部署标准化", "明确推荐 compose 文件、数据目录、备份恢复命令和生产检查清单。"],
        ["测试矩阵", "为 OAuth、Codex transform、支付回调、failover、计费边界建立固定回归用例。"],
        ["文档维护", "把本文档作为主线，附录只保留索引，避免以后又变成文件清单。"],
    ], [4.0, 11.0])

    add_heading(doc, "附录 A：关键代码索引", 1)
    add_para(doc, "下面的文件索引只作为定位入口，不再作为正文主角。真正理解系统仍然建议从上面的功能链路开始。")
    add_simple_table(doc, ["功能", "关键文件"], [
        ["前端路由与权限", "frontend/src/router/index.ts"],
        ["前端请求封装", "frontend/src/api/client.ts, frontend/src/api/auth.ts, frontend/src/api/admin/*"],
        ["前端登录状态", "frontend/src/stores/auth.ts, frontend/src/stores/app.ts"],
        ["用户页面", "frontend/src/views/user/*"],
        ["管理员页面", "frontend/src/views/admin/*"],
        ["后端启动", "backend/cmd/server/main.go, backend/cmd/server/wire_gen.go"],
        ["路由注册", "backend/internal/server/router.go, backend/internal/server/routes/*.go"],
        ["普通认证", "backend/internal/handler/auth_handler.go, backend/internal/service/auth_service.go"],
        ["OpenAI OAuth", "backend/internal/handler/admin/openai_oauth_handler.go, backend/internal/service/openai_oauth_service.go"],
        ["OpenAI/Codex 网关", "backend/internal/handler/openai_gateway_handler.go, backend/internal/service/openai_gateway_service.go"],
        ["Codex 请求转换", "backend/internal/service/openai_codex_transform.go"],
        ["账号调度", "backend/internal/service/openai_account_scheduler.go, backend/internal/service/gateway_service.go"],
        ["支付", "backend/internal/payment, backend/internal/service/payment_service.go, backend/internal/handler/payment_*"],
        ["用量计费", "backend/internal/service/usage_billing.go, backend/internal/service/usage_log.go"],
        ["数据库模型", "backend/ent/schema/*.go"],
        ["部署", "deploy/*, Dockerfile"],
    ], [4.0, 11.0])

    add_heading(doc, "附录 B：给非 TS / Go 读者的术语表", 1)
    add_simple_table(doc, ["术语", "解释"], [
        ["TypeScript", "带类型的 JavaScript，能提前发现字段拼错、返回值不匹配等问题。"],
        ["Vue", "前端页面框架，把页面拆成组件。"],
        ["Pinia", "前端全局状态管理工具，类似浏览器里的小型状态仓库。"],
        ["Axios", "前端发送 HTTP 请求的库。"],
        ["Go", "后端编程语言，适合写高并发、易部署的服务。"],
        ["Gin", "Go 的 HTTP Web 框架，负责路由和中间件。"],
        ["Ent", "Go 的数据库 ORM，通过 schema 生成类型安全的数据访问代码。"],
        ["Handler", "HTTP 处理器，负责读请求、调 service、写响应。"],
        ["Service", "业务服务，承载核心规则。"],
        ["Repository", "数据访问层，负责数据库、Redis 或外部 HTTP 客户端。"],
        ["OAuth", "授权协议，用来让系统获得第三方账号的访问凭证。"],
        ["Access Token", "短期访问令牌，用来请求上游服务。"],
        ["Refresh Token", "长期刷新令牌，用来换新的 access token。"],
        ["Failover", "失败切换，一个账号失败时换另一个账号继续尝试。"],
        ["SSE", "流式响应协议，模型边生成边返回。"],
    ], [4.0, 11.0])

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(write_manual())
